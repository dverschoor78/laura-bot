import os
import re
import json
import shutil
import hashlib
import sqlite3
import base64
import urllib.request
import anthropic
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional
from datetime import datetime
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from financeiro.lancamento import (init_db_financeiro, sugerir_categoria, CategoriaLancamento,
                                   vincular_nfe, buscar_candidatos_nfe)

load_dotenv()
TOKEN       = os.environ["TELEGRAM_BOT_TOKEN"]
DONO_ID     = int(os.environ["TELEGRAM_USER_ID"])
CLAUDE_KEY  = os.environ["CLAUDE_API_KEY"]
TEST_MODE = os.environ.get("LAURA_ENV", "prod") == "test"
DB_PATH   = Path("data/laura_test.db") if TEST_MODE else Path("data/laura.db")
UPLOADS   = Path("data/test_uploads")  if TEST_MODE else Path("data/uploads")
UPLOADS.mkdir(parents=True, exist_ok=True)

claude = anthropic.Anthropic(api_key=CLAUDE_KEY)

TIPOS = {
    "orcamento":        ("📋", "Orçamento"),
    "comprovante_pix":  ("💰", "Comprovante PIX"),
    "nota_fiscal":      ("🧾", "Nota Fiscal"),
    "foto_entrega":     ("📦", "Foto de entrega"),
    "extrato_mp":       ("🏦", "Extrato MP"),
    "nao_relacionado":  ("🗑", "Não é da obra"),
}
GGVS = ["GGV00", "GGV01", "GGV02", "GGV03"]

MESES = ["janeiro","fevereiro","março","abril","maio","junho",
         "julho","agosto","setembro","outubro","novembro","dezembro"]

DELTAD = {
    "nome":  "Verschoor Investimentos Imobiliários Ltda",
    "cnpj":  "58.358.802/0001-58",
    "ie":    "Isento",
    "end":   "Av. dos Pioneiros, 1380 – Centro – Carambeí/PR – CEP 84.145-000",
    "email": "dennis@deltad.com.br",
    "fone":  "(42) 99127-1255",
}
DELTAD_CNPJ_DIGITS = re.sub(r"\D", "", DELTAD["cnpj"])  # "58358802000158"

GGV_ENCARREGADO = {
    "GGV03": "Sabiá",
}

GGV_DESC = {
    "GGV01": "GGV01 — Matrícula 39.333, Quadra 05 Lote 02, JD das Nações, Carambeí-PR",
    "GGV02": "GGV02 — Matrícula 39.337, Quadra 05 Lote 06, JD das Nações, Carambeí-PR",
    "GGV03": "GGV03 — Matrícula 39.339, Quadra 05 Lote 08, JD das Nações, Carambeí-PR",
    "GGV00": "GGV00 — Despesas Gerais",
}

GGV_ONEDRIVE = {
    "GGV03": Path(r"C:\Users\denni\OneDrive\00 Obras\2026-06 GGV03\04 Aquisição e Execução"),
}

CONDICOES = {
    "pix_avista":  "PIX à vista",
    "pix_50_50":   "PIX 50% entrada + 50% na entrega",
}

ENDERECOS = {
    "obra_GGV01": "Rua Índia em frente ao nº139, Loteamento JD das Nações - Carambeí-PR CEP 84.145-000",
    "obra_GGV02": "Rua Índia em frente ao nº139, Loteamento JD das Nações - Carambeí-PR CEP 84.145-000",
    "obra_GGV03": "Rua Índia em frente ao nº139, Loteamento JD das Nações - Carambeí-PR CEP 84.145-000",
    "casa":        "Avenida dos Pioneiros, fundos Frederica's Coffie Huis - Carambeí-PR CEP 84.145-000",
    "escritorio":  "Avenida dos Pioneiros, 1380 - Carambeí-PR CEP 84.145-000",
    "chacara":     "Avenida dos Pioneiros, 5125 - Carambeí-PR CEP 84.145-000",
}

GGV_CODIGO_RE = re.compile(r"^\s*(GGV\d{2})\s*$", re.IGNORECASE)

class StatusPedido(str, Enum):
    A_PAGAR          = "a_pagar"
    PAGO             = "pago"
    PENDENTE_REVISAO = "pendente_revisao"
    SUBSTITUIDO      = "substituido"
    SEM_LANCAMENTO   = "sem_lancamento"

@dataclass
class Pedido:
    # Identificação
    codigo: str
    doc_id: int
    ggv:    str

    # Domínio
    status:             StatusPedido
    fornecedor:         str
    cnpj:               str
    valor_orcamento:    float
    desconto:           float
    valor_negociado:    float
    condicao_pagamento: str
    vencimento:         str
    entrega_prevista:   str

    # Datas de registro — base para construir o histórico
    doc_criado_em:  Optional[str] = None
    lanc_criado_em: Optional[str] = None
    data_pagamento: Optional[str] = None
    doc_id_nfe:            Optional[int] = None
    nfe_numero:            Optional[str] = None
    nfe_data:              Optional[str] = None
    doc_id_comprovante:    Optional[int] = None
    identificador_comprovante: Optional[str] = None
    qtd_fotos_entrega:     int = 0
    obs_entrega:           Optional[str] = None
    entregue_em:           Optional[str] = None
    categoria:             Optional[str] = None

    # Arquivos — populados por preparar_visualizacao_pedido()
    caminho_orcamento: Optional[str] = None
    caminho_docx:      Optional[str] = None
    caminho_pdf:       Optional[str] = None   # futuro

    # Histórico — populado por preparar_visualizacao_pedido()
    historico: list = field(default_factory=list)

PROMPT = """
Você recebeu um arquivo enviado para um sistema de gestão de obras de construção civil.

PASSO 1 — Classifique o documento:
[orcamento]        — cotação, orçamento, pedido de compra, lista de materiais com preços;
                     também boleto, fatura ou conta a pagar de taxa/imposto/serviço público
                     (ex: anuidade CREA, emolumentos de cartório/ONR, IPTU e taxas de prefeitura,
                     conta de energia Copel, conta de água/esgoto Sanepar)
[comprovante_pix]  — comprovante de pagamento PIX ou transferência bancária
[nota_fiscal]      — Nota Fiscal eletrônica (NF-e), DANFE, NFS-e ou recibo fiscal
[extrato_mp]       — extrato do Mercado Pago ou extrato bancário
[nao_relacionado]  — qualquer outro documento não relacionado a obras

PASSO 2 — Identifique o empreendimento (GGV):
GGV01 — Matrícula 39.333, Quadra 05 Lote 02
GGV02 — Matrícula 39.337, Quadra 05 Lote 06
GGV03 — Matrícula 39.339, Quadra 05 Lote 08
GGV00 — despesas gerais compartilhadas
nao_identificado — se não conseguir determinar

PASSO 3 — Extraia os dados em português conforme o tipo:

Se [orcamento]:
- Fornecedor:
- Ramo de atividade: (ex: Comércio de Materiais de Construção, Serralheria, Elétrica — informe como aparece no documento ou deduza pelo contexto)
- Resumo da compra: (2 a 4 palavras que identifiquem o item principal do orçamento — ex: "aço", "tubos caixa d'água", "material elétrico", "portas"; vai virar nome de arquivo)
- CNPJ/CPF:
- Chave PIX: (procure em qualquer parte do documento — dados cadastrais, cabeçalho, rodapé, condições de pagamento; pode ser CPF, CNPJ, e-mail ou telefone)
- Número do orçamento: (número ou código do orçamento emitido pelo fornecedor, se houver)
- Vendedor: (nome do vendedor ou representante que emitiu o orçamento, se houver)
- Telefone do vendedor: (telefone ou WhatsApp do vendedor, se houver)
- Itens (formato: N. Descrição do produto (QTDE UND) — R$ TOTAL; liste todos os itens do orçamento):
- Valor total:
- Desconto (valor em R$, se houver; se informado em %, calcule o valor sobre o total):
- Condição de pagamento:
- Prazo de entrega: (lead time ou data de entrega do material — NÃO a validade da proposta)
- Validade da proposta: (data até quando o preço é válido)
- Observações:

Se [comprovante_pix]:
- Data do pagamento:
- Valor:
- Favorecido:
- CNPJ/CPF do favorecido:
- Chave PIX:
- Instituição financeira:
- ID da transação: (prefira o ID EndToEnd do Pix — começa com "E", ex: E10573521202506... — se não estiver visível, use o número da transação do Mercado Pago; extraia APENAS o código, sem texto adicional)
- Identificador / Observação:

Se [nota_fiscal]:
- Número da NF:
- CNPJ/CPF do emitente:
- Nome do emitente:
- Valor total:
- Data de emissão:
- Descrição do serviço/produto: (resumo do que foi fornecido)

Se [extrato_mp]:
- Período:
- Número de transações identificadas:
- Resumo:

Se [nao_relacionado]:
- Descreva brevemente o que é o documento.

Responda EXATAMENTE neste formato (sem colchetes, sem barra, escolha um valor de cada):
TIPO:orcamento
GGV:GGV03

Valores aceitos para TIPO: orcamento, comprovante_pix, nota_fiscal, extrato_mp, nao_relacionado
Valores aceitos para GGV: GGV00, GGV01, GGV02, GGV03, nao_identificado

Em seguida, os dados extraídos conforme o tipo identificado acima.
"""

def _raiz_obra(ggv: str) -> Optional[Path]:
    """Pasta raiz da obra no OneDrive (ex: '00 Obras/2026-06 GGV03'). None se não configurada."""
    obra = buscar_obra(ggv)
    raiz = obra.get("pasta_onedrive", "")
    return Path(raiz) if raiz else None

def _pasta_pfm(ggv: str) -> Path:
    if TEST_MODE:
        pasta = Path("data/test_pfms")
    else:
        raiz  = _raiz_obra(ggv)
        pasta = (raiz / "04 Compras") if raiz else Path("data/pfms")
    pasta.mkdir(parents=True, exist_ok=True)
    return pasta

def _pasta_orcamentos(ggv: str) -> Path:
    pasta = _pasta_pfm(ggv) / "00 Orçamentos"
    pasta.mkdir(parents=True, exist_ok=True)
    return pasta

def _pasta_controle_financeiro(ggv: str) -> Path:
    if TEST_MODE:
        pasta = Path("data/test_pfms") / "01 Controle financeiro"
    else:
        raiz  = _raiz_obra(ggv)
        pasta = (raiz / "01 Controle financeiro") if raiz else Path("data/pfms") / "01 Controle financeiro"
    pasta.mkdir(parents=True, exist_ok=True)
    return pasta

def _pasta_entrega(ggv: str) -> Path:
    if TEST_MODE:
        pasta = Path("data/test_pfms") / "05 Entrega"
    else:
        raiz  = _raiz_obra(ggv)
        pasta = (raiz / "05 Entrega") if raiz else Path("data/pfms") / "05 Entrega"
    pasta.mkdir(parents=True, exist_ok=True)
    return pasta

def _nome_arquivo_seguro(s: str, max_len: int = 60) -> str:
    """Remove caracteres inválidos em nome de arquivo do Windows. Vazio/"A PREENCHER" -> ""."""
    if not s or s == "A PREENCHER":
        return ""
    s = re.sub(r'[<>:"/\\|?*]', "", s).strip()
    s = re.sub(r"\s+", " ", s)
    return s[:max_len].strip()

def _nome_base_pfm(pfm_codigo: str, fornecedor: str, resumo: str, prefixo: str = "") -> str:
    """Monta o nome de arquivo do PFM: 'GGV03-008 - Fornecedor - Resumo'."""
    partes = [f"{prefixo}{pfm_codigo}"]
    for campo in (fornecedor, resumo):
        seguro = _nome_arquivo_seguro(campo)
        if seguro:
            partes.append(seguro)
    return " - ".join(partes)

def _data_para_arquivo(data_str: str) -> str:
    """Converte data extraída (DD/MM/AAAA ou "DD de mês de AAAA") para AAAA-MM-DD. Usa hoje se não conseguir."""
    if data_str and data_str != "A PREENCHER":
        m = re.match(r"(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})", data_str)
        if m:
            d, mth, y = m.groups()
            if len(y) == 2:
                y = "20" + y
            try:
                return f"{y}-{int(mth):02d}-{int(d):02d}"
            except ValueError:
                pass
        m = re.search(r"(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})", data_str, re.IGNORECASE)
        if m:
            d, mes_nome, y = m.groups()
            if mes_nome.lower() in MESES:
                mth = MESES.index(mes_nome.lower()) + 1
                try:
                    return f"{y}-{mth:02d}-{int(d):02d}"
                except ValueError:
                    pass
    return datetime.now().strftime("%Y-%m-%d")

def _arquivar_documento(pfm_codigo: str, sufixo: str, caminho_original, data_str, pasta_fn):
    """Copia um documento vinculado a um pedido para a pasta certa, nome padronizado. Falha silenciosa."""
    if not caminho_original or not Path(caminho_original).exists():
        return
    ggv = pfm_codigo.split("-")[0]
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute("SELECT fornecedor FROM lancamentos WHERE pfm_codigo=?", (pfm_codigo,)).fetchone()
    fornecedor   = row[0] if row else ""
    data_arquivo = _data_para_arquivo(data_str)
    nome = f"{data_arquivo} {pfm_codigo}"
    forn_seguro = _nome_arquivo_seguro(fornecedor)
    if forn_seguro:
        nome += f" {forn_seguro}"
    nome += f" - {sufixo}"
    ext = Path(caminho_original).suffix
    destino = pasta_fn(ggv) / f"{nome}{ext}"
    try:
        shutil.copy2(caminho_original, destino)
    except OSError:
        pass

def _arquivar_doc_financeiro(pfm_codigo: str, sufixo: str, caminho_original, data_str: str):
    """Copia comprovante/NF-e para '01 Controle financeiro', nome padronizado."""
    _arquivar_documento(pfm_codigo, sufixo, caminho_original, data_str, _pasta_controle_financeiro)

# ── Banco ──────────────────────────────────────────────────────────────────

def _migrar_obras(con):
    """Pré-popula a tabela obras a partir dos dados que eram hardcoded. Idempotente."""
    dados = [
        ("GGV00", "Despesas Gerais", "", "", "",
         "Dennis Verschoor", "(42) 99127-1255", "", 1),
        ("GGV01", "Matrícula 39.333, Quadra 05 Lote 02, JD das Nações, Carambeí-PR",
         "Rua Índia em frente ao nº139, JD das Nações - Carambeí-PR CEP 84.145-000",
         "", "", "Dennis Verschoor", "(42) 99127-1255", "", 1),
        ("GGV02", "Matrícula 39.337, Quadra 05 Lote 06, JD das Nações, Carambeí-PR",
         "Rua Índia em frente ao nº139, JD das Nações - Carambeí-PR CEP 84.145-000",
         "", "", "Dennis Verschoor", "(42) 99127-1255", "", 1),
        ("GGV03", "Matrícula 39.339, Quadra 05 Lote 08, JD das Nações, Carambeí-PR",
         "Rua Índia em frente ao nº139, JD das Nações - Carambeí-PR CEP 84.145-000",
         "Sabiá", "(42) 98439-9498",
         "Dennis Verschoor", "(42) 99127-1255",
         r"C:\Users\denni\OneDrive\00 Obras\2026-06 GGV03\04 Aquisição e Execução", 1),
    ]
    for d in dados:
        con.execute("""
            INSERT OR IGNORE INTO obras
            (codigo, descricao, endereco_entrega, encarregado_nome, encarregado_fone,
             responsavel_nome, responsavel_fone, pasta_onedrive, ativa)
            VALUES (?,?,?,?,?,?,?,?,?)
        """, d)

def init_db():
    with sqlite3.connect(DB_PATH) as con:
        con.execute("""
            CREATE TABLE IF NOT EXISTS documentos (
                id               INTEGER PRIMARY KEY AUTOINCREMENT,
                nome             TEXT,
                caminho          TEXT,
                hash             TEXT UNIQUE,
                tipo             TEXT DEFAULT 'desconhecido',
                ggv              TEXT DEFAULT 'nao_identificado',
                dados_claude     TEXT,
                condicao_pgto    TEXT,
                data_entrega     TEXT,
                endereco_entrega TEXT,
                desconto_rs      TEXT,
                pfm_numero       INTEGER,
                status           TEXT DEFAULT 'recebido',
                criado_em        TEXT DEFAULT (datetime('now','localtime'))
            )
        """)
        for col in ["tipo", "ggv", "dados_claude", "condicao_pgto", "data_entrega", "endereco_entrega", "desconto_rs TEXT", "pfm_numero INTEGER", "vencimento_pgto TEXT", "encarregado TEXT", "rev_numero INTEGER DEFAULT 0", "caminho_pfm TEXT"]:
            try:
                con.execute(f"ALTER TABLE documentos ADD COLUMN {col}")
            except Exception:
                pass
        con.execute("""
            CREATE TABLE IF NOT EXISTS lancamentos (
                id                    INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id                INTEGER NOT NULL,
                pfm_codigo            TEXT NOT NULL UNIQUE,
                ggv                   TEXT,
                fornecedor            TEXT,
                valor                 REAL,
                data_prevista_entrega TEXT,
                vencimento_pagamento  TEXT,
                status                TEXT DEFAULT 'a_pagar',
                criado_em             TEXT DEFAULT (datetime('now','localtime'))
            )
        """)
        for col in ["valor_pago REAL", "data_pagamento TEXT", "doc_id_comprovante INTEGER",
                    "identificador_comprovante TEXT", "doc_id_entrega INTEGER",
                    "obs_entrega TEXT", "entregue_em TEXT"]:
            try:
                con.execute(f"ALTER TABLE lancamentos ADD COLUMN {col}")
            except Exception:
                pass
        con.execute("""
            CREATE TABLE IF NOT EXISTS entrega_fotos (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                pfm_codigo  TEXT NOT NULL,
                doc_id      INTEGER NOT NULL,
                legenda     TEXT,
                criado_em   TEXT DEFAULT (datetime('now','localtime'))
            )
        """)
        con.execute("""
            CREATE TABLE IF NOT EXISTS obras (
                codigo            TEXT PRIMARY KEY,
                descricao         TEXT,
                endereco_entrega  TEXT,
                encarregado_nome  TEXT,
                encarregado_fone  TEXT,
                responsavel_nome  TEXT,
                responsavel_fone  TEXT,
                pasta_onedrive    TEXT,
                ativa             INTEGER DEFAULT 1,
                criado_em         TEXT DEFAULT (datetime('now','localtime'))
            )
        """)
        _migrar_obras(con)
        con.execute("""
            CREATE TABLE IF NOT EXISTS fornecedores (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                nome         TEXT NOT NULL,
                razao_social TEXT,
                cnpj         TEXT,
                cpf          TEXT,
                chave_pix    TEXT,
                email        TEXT,
                whatsapp     TEXT,
                contato      TEXT,
                logradouro   TEXT,
                numero       TEXT,
                bairro       TEXT,
                cidade       TEXT,
                uf           TEXT,
                cep          TEXT,
                origem       TEXT,
                criado_em    TEXT DEFAULT (datetime('now','localtime')),
                UNIQUE(cnpj),
                UNIQUE(cpf)
            )
        """)
        for col in ["ramo TEXT", "receita_pendente INTEGER DEFAULT 0"]:
            try:
                con.execute(f"ALTER TABLE fornecedores ADD COLUMN {col}")
            except Exception:
                pass
    init_db_financeiro(DB_PATH)

def buscar_obra(codigo):
    """Retorna dict com dados da obra ou {} se não encontrada."""
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(
            "SELECT codigo, descricao, endereco_entrega, encarregado_nome, encarregado_fone, "
            "responsavel_nome, responsavel_fone, pasta_onedrive FROM obras WHERE codigo=?",
            (codigo,)
        ).fetchone()
    if not row:
        return {}
    keys = ["codigo", "descricao", "endereco_entrega", "encarregado_nome",
            "encarregado_fone", "responsavel_nome", "responsavel_fone", "pasta_onedrive"]
    return dict(zip(keys, row))

def atualizar_obra(codigo, **kwargs):
    cols = ", ".join(f"{k}=?" for k in kwargs)
    vals = list(kwargs.values()) + [codigo]
    with sqlite3.connect(DB_PATH) as con:
        con.execute(f"UPDATE obras SET {cols} WHERE codigo=?", vals)

def criar_obra(codigo, descricao=""):
    """Insere nova obra. Retorna True se criada, False se já existia."""
    with sqlite3.connect(DB_PATH) as con:
        cur = con.execute(
            "INSERT OR IGNORE INTO obras (codigo, descricao, ativa) VALUES (?, ?, 1)",
            (codigo.upper(), descricao)
        )
        return cur.rowcount == 1

def ja_existe(hash_arquivo):
    if TEST_MODE:
        return None
    with sqlite3.connect(DB_PATH) as con:
        return con.execute("SELECT id FROM documentos WHERE hash = ?", (hash_arquivo,)).fetchone()

def registrar(nome, caminho, hash_arquivo, tipo, ggv, dados_claude):
    with sqlite3.connect(DB_PATH) as con:
        cur = con.execute(
            "INSERT INTO documentos (nome, caminho, hash, tipo, ggv, dados_claude) VALUES (?,?,?,?,?,?)",
            (nome, str(caminho), hash_arquivo, tipo, ggv, dados_claude)
        )
        return cur.lastrowid

def atualizar(doc_id, **campos):
    sets = ", ".join(f"{k}=?" for k in campos)
    with sqlite3.connect(DB_PATH) as con:
        con.execute(f"UPDATE documentos SET {sets} WHERE id=?", (*campos.values(), doc_id))

def registrar_lancamento(doc_id, pfm_codigo, ggv, fornecedor, valor_v, data_entrega, categoria=None):
    """Insere lançamento A PAGAR. Idempotente: se pfm_codigo já existe, retorna o existente."""
    forn_ok  = fornecedor and fornecedor != "A PREENCHER"
    valor_ok = valor_v and valor_v > 0
    status   = "a_pagar" if (forn_ok and valor_ok) else "pendente_revisao"
    cat_val  = categoria.value if categoria else None
    with sqlite3.connect(DB_PATH) as con:
        cur = con.execute(
            """INSERT OR IGNORE INTO lancamentos
               (doc_id, pfm_codigo, ggv, fornecedor, valor, data_prevista_entrega, status, categoria)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (doc_id, pfm_codigo, ggv, fornecedor, valor_v, data_entrega, status, cat_val)
        )
        ja_existia = cur.rowcount == 0
        if ja_existia:
            row = con.execute(
                "SELECT status FROM lancamentos WHERE pfm_codigo=?", (pfm_codigo,)
            ).fetchone()
            status = row[0] if row else status
    return status, ja_existia

def _dados_doc(doc_id):
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute("SELECT dados_claude FROM documentos WHERE id=?", (doc_id,)).fetchone()
    return row[0] if row else ""

_CAMPOS_RESUMO_RE = re.compile(
    r"^\s*-\s*\*{0,2}(Desconto|Condição de pagamento|Prazo de entrega|Data.prazo de entrega)\b",
    re.IGNORECASE
)

PFM_CODIGO_RE = re.compile(r"\b(GGV\d{2}-\d{3}(?:-R\d+)?)\b", re.IGNORECASE)

def _dados_display(dados):
    """Remove do texto do Claude os campos que já aparecem no bloco de resumo."""
    linhas = [l for l in dados.split('\n') if not _CAMPOS_RESUMO_RE.match(l)]
    return '\n'.join(linhas).strip()

def _resumo_gerar(doc_id):
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(
            "SELECT dados_claude, tipo, ggv, condicao_pgto, data_entrega, desconto_rs, "
            "vencimento_pgto, encarregado, endereco_entrega FROM documentos WHERE id=?",
            (doc_id,)
        ).fetchone()
    if not row:
        return "Documento não encontrado.", None
    dados, tipo, ggv, condicao, data_ent, desconto_rs, vencimento, encarregado, endereco = row

    label_ggv = f"Obra {ggv}" if ggv != "nao_identificado" else "Obra não identificada"

    def _v(val):
        return None if (not val or val == "A PREENCHER") else val

    fornecedor = _v(_campo(dados, "Fornecedor")) or "Fornecedor não identificado"
    cnpj       = _v(_campo(dados, "CNPJ/CPF"))
    pix        = _v(_campo(dados, "Chave PIX"))
    vendedor   = _v(_campo(dados, "Vendedor"))
    vend_fone  = _v(_campo(dados, "Telefone do vendedor"))
    cond       = _v(condicao) or _v(_campo(dados, "Condição de pagamento"))
    entrega    = _v(data_ent) or _v(_campo(dados, "Prazo de entrega"))
    validade   = _v(_campo(dados, "Validade da proposta"))
    obs        = _obs(dados).strip()
    _obra      = buscar_obra(ggv)
    enc        = _v(encarregado) or _obra.get("encarregado_nome")

    subtotal_v, desconto_v, total_final_v = _calcular_totais(dados, desconto_rs)

    SEP = "──────────────────────────────────"
    linhas = []

    # Bloco 1 — Obra
    linhas.append(f"<b>{_esc_html(label_ggv)}</b>")
    linhas.append(SEP)

    # Bloco 2 — Fornecedor + dados de pagamento do fornecedor
    linhas.append(_esc_html(fornecedor))
    if cnpj: linhas.append(f"CNPJ  {_esc_html(cnpj)}")
    if pix:  linhas.append(f"PIX   {_esc_html(pix)}")
    if vendedor:
        cont = _esc_html(vendedor)
        if vend_fone:
            cont += f"  {_esc_html(vend_fone)}"
        linhas.append(f"Contato   {cont}")
    linhas.append(SEP)

    # Bloco 3 — Itens + Total
    bloco_itens = _bloco_itens(dados)
    for linha in bloco_itens.splitlines():
        if linha.strip():
            linhas.append(_esc_html(linha.strip()))
    if subtotal_v > 0:
        linhas.append(f"Total — R$ {_fmt_brl(subtotal_v)}")
    linhas.append(SEP)

    # Bloco 4 — Financeiro (desconto → valor final → condição → vencimento)
    if desconto_v > 0 and subtotal_v > 0:
        pct = desconto_v / subtotal_v * 100
        linhas.append(f"Desconto — R$ {_fmt_brl(desconto_v)} ({pct:.0f}%)")
    if total_final_v > 0:
        linhas.append(f"<b>Valor final — R$ {_fmt_brl(total_final_v)}</b>")
    else:
        linhas.append("<b>Valor final: não informado</b>")
    linhas.append(_esc_html(cond) if cond else "Pagamento: não informado")
    linhas.append(_esc_html(_v(vencimento)) if _v(vencimento) else "Vencimento: não informado")
    linhas.append(SEP)

    # Bloco 5 — Logística
    linhas.append(f"Entrega: {_esc_html(entrega) if entrega else 'não informada'}")
    linhas.append(f"Endereço: {_esc_html(endereco) if endereco else 'não informado'}")
    if validade:
        linhas.append(f"Válido até: {_esc_html(validade)}")
    if enc:
        enc_fone = _obra.get("encarregado_fone", "")
        enc_txt  = f"{_esc_html(enc)} {enc_fone}".strip() if enc_fone else _esc_html(enc)
        linhas.append(f"Dúvidas: Dennis {DELTAD['fone']} ou {enc_txt}, encarregado")
    else:
        linhas.append(f"Dúvidas: Dennis {DELTAD['fone']}")
    linhas.append(SEP)

    # Bloco 6 — Observações (sempre mostrado)
    linhas.append(f"Obs: {_esc_html(obs) if obs else 'não informado'}")

    # Atenção quando obra não definida
    if ggv == "nao_identificado":
        linhas.append(SEP)
        linhas.append("⚠️ Defina a obra antes de gerar o Pedido de Compra.")

    return "\n".join(linhas), teclado_orcamento(doc_id, tipo, ggv)

_FORN_COLS = ["nome", "razao_social", "cnpj", "cpf", "chave_pix", "email",
              "whatsapp", "logradouro", "numero", "bairro", "cidade", "uf", "cep", "ramo"]

def _consultar_receita(cnpj_digits: str, timeout: float = 4.0) -> Optional[dict]:
    """Consulta CNPJ na Receita Federal via BrasilAPI. Nunca levanta — retorna None em qualquer falha."""
    try:
        req = urllib.request.Request(
            f"https://brasilapi.com.br/api/cnpj/v1/{cnpj_digits}",
            headers={"User-Agent": "laura-bot"}
        )
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return {
            "razao_social": data.get("razao_social") or None,
            "cidade":       (data.get("municipio") or "").title() or None,
            "uf":           data.get("uf") or None,
        }
    except Exception:
        return None

def _criar_fornecedor_auto(nome_claude, cnpj_claude, ramo_claude, doc_id):
    """Cadastra um fornecedor novo a partir de um orçamento com CNPJ ainda não conhecido.
    Tenta enriquecer com dado oficial da Receita; se a consulta falhar, marca para sincronizar depois."""
    if not cnpj_claude or cnpj_claude == "A PREENCHER":
        return
    cnpj_digits = re.sub(r"\D", "", cnpj_claude)
    if len(cnpj_digits) != 14 or cnpj_digits == DELTAD_CNPJ_DIGITS:
        return
    receita = _consultar_receita(cnpj_digits)
    ramo = ramo_claude if ramo_claude and ramo_claude != "A PREENCHER" else None
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            "INSERT OR IGNORE INTO fornecedores "
            "(nome, cnpj, razao_social, cidade, uf, ramo, origem, receita_pendente) "
            "VALUES (?,?,?,?,?,?,?,?)",
            (nome_claude, cnpj_claude,
             receita["razao_social"] if receita else None,
             receita["cidade"] if receita else None,
             receita["uf"] if receita else None,
             ramo, f"Cadastro automático — doc {doc_id}",
             0 if receita else 1)
        )

def buscar_fornecedor(nome_claude, cnpj_claude=None):
    """Busca no BD: 1º por CNPJ exato, 2º por prefixo do nome."""
    with sqlite3.connect(DB_PATH) as con:
        sel = f"SELECT {', '.join(_FORN_COLS)} FROM fornecedores"

        # 1. CNPJ — mais confiável; ignora o nosso próprio CNPJ (dado para fatura extraído errado)
        if cnpj_claude and cnpj_claude != "A PREENCHER":
            cnpj_digits = re.sub(r"\D", "", cnpj_claude)
            if cnpj_digits != DELTAD_CNPJ_DIGITS:
                row = con.execute(
                    f"{sel} WHERE REPLACE(REPLACE(REPLACE(cnpj,'.','' ),'/',''),'-','') = ? LIMIT 1",
                    (cnpj_digits,)
                ).fetchone()
                if row:
                    return dict(zip(_FORN_COLS, row))

        # 2. Prefixo do primeiro token do nome (evita falsos positivos de substring)
        if nome_claude and nome_claude != "A PREENCHER":
            primeiro = nome_claude.strip().upper().split()[0]
            row = con.execute(
                f"{sel} WHERE UPPER(nome) LIKE ? OR UPPER(razao_social) LIKE ? LIMIT 1",
                (f"{primeiro}%", f"{primeiro}%")
            ).fetchone()
            if row:
                return dict(zip(_FORN_COLS, row))

    return None

# ── PFM ───────────────────────────────────────────────────────────────────

def _esc_html(s):
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _campo(dados, nome):
    nao_encontrado = {
        "não identificado", "nao identificado",
        "não especificado", "nao especificado",
        "não informado", "nao informado",
        "n/a", "—", "-", "",
    }
    for linha in dados.splitlines():
        stripped = linha.strip().lstrip("- *")
        if stripped.lower().startswith(nome.lower() + ":"):
            val = stripped.split(":", 1)[1].strip().strip("*").strip()
            if val.lower() not in nao_encontrado:
                return val
    return "A PREENCHER"

def _substituir_campo(dados, nome, novo_valor):
    linhas = dados.splitlines()
    for i, linha in enumerate(linhas):
        stripped = linha.strip().lstrip("- *")
        if stripped.lower().startswith(nome.lower() + ":"):
            linhas[i] = f"{nome}: {novo_valor}"
            return "\n".join(linhas)
    return dados + f"\n{nome}: {novo_valor}"

def _bloco_itens(dados):
    linhas = dados.splitlines()
    capturando, resultado = False, []
    for linha in linhas:
        stripped = linha.strip().lstrip("- *")
        if re.match(r"^(itens|materiais)", stripped, re.IGNORECASE) and ":" in stripped:
            capturando = True
            continue
        if capturando:
            if stripped and not re.match(r"^\d+\.", stripped) and ":" in stripped and not stripped.startswith("http"):
                break
            resultado.append(linha)
    return "\n".join(resultado).strip() or "Nenhum item encontrado."

def _substituir_itens(dados, novo_bloco):
    linhas = dados.splitlines()
    inicio, fim, capturando = None, len(linhas), False
    for i, linha in enumerate(linhas):
        stripped = linha.strip().lstrip("- *")
        if re.match(r"^(itens|materiais)", stripped, re.IGNORECASE) and ":" in stripped:
            inicio = i
            capturando = True
            continue
        if capturando:
            if stripped and not re.match(r"^\d+\.", stripped) and ":" in stripped and not stripped.startswith("http"):
                fim = i
                break
    if inicio is None:
        return dados + f"\nItens:\n{novo_bloco}"
    return "\n".join(linhas[:inicio + 1] + novo_bloco.splitlines() + linhas[fim:])

def _recalcular_itens(dados: str) -> str:
    """Após edição de itens, recalcula total de cada linha (qtde × unit) e atualiza Valor total."""
    linhas_out = []
    capturando = False
    novo_total = 0.0
    for linha in dados.splitlines():
        stripped = linha.strip().lstrip("- *")
        if re.match(r"^(itens|materiais)", stripped, re.IGNORECASE) and ":" in stripped:
            capturando = True
            linhas_out.append(linha)
            continue
        if capturando:
            if re.match(r"^(valor total|condição|condicao|prazo|validade|observ)", stripped, re.IGNORECASE):
                capturando = False
                linhas_out.append(linha)
                continue
            m = ITEM_RE.match(stripped)
            if m:
                desc, qtde_str, und, val1, val2 = m.groups()
                qtde_v = _parse_brl(qtde_str)
                num_m = re.match(r"^(\d+)\.", stripped)
                num = num_m.group(1) if num_m else "?"
                if val2:
                    unit_v  = _parse_brl(val1)
                    total_v = round(qtde_v * unit_v, 2)
                    novo_total += total_v
                    linhas_out.append(
                        f"{num}. {desc.strip()} ({qtde_str} {und.upper()}) — R$ {_fmt_brl(unit_v)} cada = R$ {_fmt_brl(total_v)}"
                    )
                    continue
                else:
                    novo_total += _parse_brl(val1)
            linhas_out.append(linha)
            continue
        linhas_out.append(linha)
    resultado = "\n".join(linhas_out)
    if novo_total > 0:
        resultado = _substituir_campo(resultado, "Valor total", f"R$ {_fmt_brl(novo_total)}")
    return resultado

ITEM_RE = re.compile(
    r"^\d+\.\s+(.+?)\s+\(([0-9,.]+)\s+([A-Za-z]{1,4})\)\s*[—–\-]+\s*R\$\s*([0-9.,]+)"
    r"(?:\s*cada\s*=\s*R\$\s*([0-9.,]+))?",
    re.IGNORECASE,
)

def _parse_brl(s):
    s = s.strip().replace(" ", "")
    if "," in s:
        return float(s.replace(".", "").replace(",", "."))
    return float(s.replace(",", "."))

def _fmt_brl(v):
    s = f"{v:,.2f}"                                        # "6,292.93"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")  # "6.292,93"

def _itens(dados):
    """Retorna lista de dicts {desc, und, qtde, unit, total, _total_v} ou strings como fallback."""
    resultado, capturando = [], False
    for linha in dados.splitlines():
        stripped = linha.strip().lstrip("- *")
        if re.match(r"^(itens|materiais)", stripped, re.IGNORECASE) and ":" in stripped:
            capturando = True
            continue
        if capturando:
            if not stripped:
                continue
            if re.match(r"^(valor total|condição|condicao|prazo|validade|observ)", stripped, re.IGNORECASE):
                break
            m = ITEM_RE.match(stripped)
            if m:
                desc, qtde_str, und, val1, val2 = m.groups()
                qtde_v = _parse_brl(qtde_str)
                if val2:
                    unit_v  = _parse_brl(val1)
                    total_v = round(qtde_v * unit_v, 2)
                else:
                    total_v = _parse_brl(val1)
                    unit_v  = total_v / qtde_v if qtde_v else 0
                resultado.append({
                    "desc":     desc.strip(),
                    "und":      und.upper(),
                    "qtde":     qtde_str,
                    "unit":     _fmt_brl(unit_v),
                    "total":    _fmt_brl(total_v),
                    "_total_v": total_v,
                })
            elif stripped:
                resultado.append(stripped)
    return resultado

def _obs(dados):
    resultado, capturando = [], False
    for linha in dados.splitlines():
        stripped = linha.strip()
        if stripped.lower().startswith("observaç"):
            capturando = True
            continue
        if capturando and stripped:
            resultado.append(stripped.lstrip("- *"))
    return "\n".join(resultado)

def _calcular_totais(dados, desconto_rs):
    """Retorna (subtotal_v, desconto_v, total_final_v) a partir dos dados extraídos."""
    itens = _itens(dados)
    subtotal_v = sum(i.get("_total_v", 0) for i in itens if isinstance(i, dict))
    if not subtotal_v:
        try:
            subtotal_v = _parse_brl(re.sub(r"[^\d,.]", "", _campo(dados, "Valor total")))
        except Exception:
            subtotal_v = 0.0
    desconto_v = 0.0
    desconto_raw = desconto_rs or (
        _campo(dados, "Desconto") if _campo(dados, "Desconto") != "A PREENCHER" else None
    )
    if desconto_raw:
        try:
            desconto_v = _parse_brl(re.sub(r"[^\d,.]", "", str(desconto_raw)))
        except Exception:
            desconto_v = 0.0
    return subtotal_v, desconto_v, subtotal_v - desconto_v

def _data_extenso(dt):
    return f"Carambeí, {dt.day} de {MESES[dt.month-1]} de {dt.year}."

def _cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_col_widths(tbl, widths_cm):
    tbl.autofit = False
    for i, w in enumerate(widths_cm):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(w)

def _secao_row(tbl, texto, ncols, bg="1F3864"):
    """Linha de cabeçalho de seção (fundo azul, texto branco)."""
    row = tbl.add_row()
    cell = row.cells[0]
    for i in range(1, ncols):
        cell.merge(row.cells[i])
    _cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

def _kv_row(tbl, label, valor, label_bg="D9D9D9"):
    """Linha label | valor em 2 colunas."""
    row = tbl.add_row()
    c_l = row.cells[0]
    _cell_bg(c_l, label_bg)
    r = c_l.paragraphs[0].add_run(label)
    r.bold = True
    r.font.size = Pt(8)
    c_r = row.cells[1]
    c_r.paragraphs[0].add_run(str(valor or "")).font.size = Pt(9)
    return row

def proximo_pfm_numero(ggv):
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(
            "SELECT MAX(pfm_numero) FROM documentos WHERE ggv=? AND pfm_numero IS NOT NULL",
            (ggv,)
        ).fetchone()
    return (row[0] or 0) + 1

_PC_CSS = """
*, *::before, *::after { margin: 0; padding: 0; box-sizing: border-box; }
@page { size: A4; margin: 0; }
body { font-family: 'Inter', system-ui, -apple-system, sans-serif; -webkit-font-smoothing: antialiased; }
.page { background: #fff; width: 210mm; min-height: 297mm; padding: 16mm 18mm 12mm; display: flex; flex-direction: column; }
.header { display: flex; justify-content: space-between; align-items: flex-start; }
.company-brand { font-size: 12px; font-weight: 600; letter-spacing: 0.01em; color: #111827; }
.company-meta { font-size: 9.5px; color: #9CA3AF; margin-top: 6px; line-height: 1.75; }
.doc-meta { text-align: right; }
.doc-tipo { font-size: 9px; font-weight: 500; letter-spacing: 0.14em; text-transform: uppercase; color: #9CA3AF; }
.doc-number { font-size: 27px; font-weight: 700; color: #111827; letter-spacing: -0.025em; line-height: 1; margin-top: 5px; }
.doc-date { font-size: 9.5px; color: #9CA3AF; margin-top: 7px; }
.rule { border: none; border-top: 1px solid #E5E7EB; }
.rule-gap { margin: 28px 0; }
.context-block { display: grid; grid-template-columns: 1fr 1fr; gap: 20px 48px; }
.ctx-label { font-size: 9px; font-weight: 600; letter-spacing: 0.15em; text-transform: uppercase; color: #9CA3AF; margin-bottom: 5px; }
.ctx-value { font-size: 10.5px; color: #6B7280; line-height: 1.7; }
.supplier { margin-top: 6px; }
.supplier-name { font-size: 22px; font-weight: 600; color: #111827; letter-spacing: -0.015em; line-height: 1.2; }
.supplier-trade { font-size: 11px; color: #6B7280; margin-top: 5px; }
.supplier-meta { font-size: 10px; color: #9CA3AF; margin-top: 3px; }
.section-label { font-size: 9px; font-weight: 600; letter-spacing: 0.15em; text-transform: uppercase; color: #9CA3AF; margin-bottom: 16px; }
.item { display: flex; justify-content: space-between; align-items: flex-start; padding: 13px 0; border-bottom: 1px solid #F9FAFB; }
.item:last-child { border-bottom: none; }
.item-left { display: flex; gap: 16px; }
.item-num { font-size: 10px; font-weight: 500; color: #D1D5DB; min-width: 20px; padding-top: 1px; flex-shrink: 0; }
.item-desc { font-size: 11.5px; font-weight: 500; color: #374151; line-height: 1.4; }
.item-qty { font-size: 10px; color: #9CA3AF; margin-top: 3px; }
.item-value { font-size: 11.5px; font-weight: 500; color: #374151; white-space: nowrap; }
.financial-outer { display: flex; justify-content: flex-end; margin-top: 24px; }
.financial-inner { min-width: 216px; }
.fin-row { display: flex; justify-content: space-between; gap: 48px; padding: 4px 0; }
.fin-l { font-size: 10px; color: #9CA3AF; }
.fin-v { font-size: 10px; color: #6B7280; }
.fin-rule { border: none; border-top: 1px solid #E5E7EB; margin: 10px 0; }
.fin-total-row { display: flex; justify-content: space-between; align-items: baseline; gap: 48px; }
.fin-total-l { font-size: 10px; font-weight: 700; letter-spacing: 0.1em; text-transform: uppercase; color: #111827; }
.fin-total-v { font-size: 20px; font-weight: 700; color: #111827; letter-spacing: -0.02em; }
.bottom { display: grid; grid-template-columns: 1fr 1fr; gap: 48px; }
.bottom-label { font-size: 9px; font-weight: 600; letter-spacing: 0.15em; text-transform: uppercase; color: #9CA3AF; margin-bottom: 10px; }
.bottom-main { font-size: 11.5px; font-weight: 500; color: #374151; margin-bottom: 5px; }
.bottom-detail { font-size: 10px; color: #6B7280; line-height: 1.7; }
.footer-tagline { margin-top: auto; padding-top: 16px; text-align: center; font-size: 11px; font-style: italic; color: #B8BFC9; letter-spacing: 0.01em; }
"""

def _gerar_html_pc(doc_id: int) -> str:
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(
            "SELECT ggv, dados_claude, condicao_pgto, data_entrega, endereco_entrega, "
            "desconto_rs, vencimento_pgto, criado_em FROM documentos WHERE id=?",
            (doc_id,)
        ).fetchone()
        if row is None:
            raise ValueError(f"Documento {doc_id} não encontrado.")
        ggv, dados, condicao, data_ent_db, end_db, desconto_rs, vencimento, criado_em = row
        pfm_row = con.execute(
            "SELECT pfm_codigo FROM lancamentos WHERE doc_id=? ORDER BY id DESC LIMIT 1",
            (doc_id,)
        ).fetchone()
    pfm_codigo  = pfm_row[0] if pfm_row else "—"
    obra        = buscar_obra(ggv)

    nome_claude = _campo(dados, "Fornecedor")
    cnpj_claude = _campo(dados, "CNPJ/CPF")
    forn_db     = buscar_fornecedor(nome_claude, cnpj_claude)
    if forn_db:
        fornecedor = forn_db.get("razao_social") or forn_db.get("nome") or nome_claude
        cnpj       = forn_db.get("cnpj") or forn_db.get("cpf") or cnpj_claude
        ramo       = forn_db.get("ramo") or _campo(dados, "Ramo de atividade")
        _cidade    = forn_db.get("cidade") or ""
        _uf        = forn_db.get("uf") or ""
        if len(_cidade) > 30 or "/" in _cidade or any(c.isdigit() for c in _cidade):
            _cidade = _uf = ""
        forn_local = " · ".join(filter(None, [_cidade, _uf]))
        pix        = forn_db.get("chave_pix") or _campo(dados, "Chave PIX")
    else:
        fornecedor = nome_claude
        cnpj       = cnpj_claude
        ramo       = _campo(dados, "Ramo de atividade")
        forn_local = ""
        pix        = _campo(dados, "Chave PIX")

    nr_orc    = _campo(dados, "Número do orçamento")
    vendedor  = _campo(dados, "Vendedor")
    vend_fone = _campo(dados, "Telefone do vendedor")

    itens      = _itens(dados)
    subtotal_v, desconto_v, total_final_v = _calcular_totais(dados, desconto_rs)
    desconto_pct = (desconto_v / subtotal_v * 100) if subtotal_v > 0 and desconto_v > 0 else 0

    now          = datetime.now()
    data_emissao = f"{now.day} de {MESES[now.month-1]} de {now.year}"

    def _h(s): return _esc_html(str(s)) if s and s != "A PREENCHER" else ""

    # Bloco Origem
    origem_linhas = []
    if _h(nr_orc):
        origem_linhas.append(f"Orçamento #{_h(nr_orc)}")
    resp_nome = obra.get("responsavel_nome") or "Dennis"
    resp_fone = obra.get("responsavel_fone") or DELTAD["fone"]
    origem_linhas.append("Negociado via WhatsApp")
    contatos = [f"{_h(resp_nome)} {_h(resp_fone)}".strip()]
    if _h(vendedor):
        v = _h(vendedor)
        if _h(vend_fone):
            v += f" {_h(vend_fone)}"
        contatos.append(v)
    origem_linhas.append(" &middot; ".join(contatos))
    origem_html = "<br>".join(origem_linhas)

    # Bloco Entrega
    enc_nome = obra.get("encarregado_nome", "")
    enc_fone = obra.get("encarregado_fone", "")
    entrega_linhas = [f"Obra {ggv}"]
    if _h(data_ent_db):
        entrega_linhas.append(f"Até {_h(data_ent_db)}")
    if enc_nome:
        enc_str = f"Encarregado: {_h(enc_nome)}"
        if enc_fone:
            enc_str += f" {_h(enc_fone)}"
        entrega_linhas.append(enc_str)
    entrega_html = "<br>".join(entrega_linhas)

    # Itens
    items_html = ""
    for i, item in enumerate(itens, 1):
        if not isinstance(item, dict):
            items_html += f'<div class="item"><div class="item-left"><span class="item-num">{i:02d}</span><div><div class="item-desc">{_esc_html(str(item))}</div></div></div></div>'
            continue
        und      = _esc_html(item.get("und", "un"))
        qty_line = f'<div class="item-qty">{_esc_html(item["qtde"])} {und} &nbsp;&middot;&nbsp; R$ {_esc_html(item["unit"])}/{und}</div>' if item.get("qtde") else ""
        items_html += (
            f'<div class="item">'
            f'<div class="item-left"><span class="item-num">{i:02d}</span>'
            f'<div><div class="item-desc">{_esc_html(item.get("desc",""))}</div>{qty_line}</div></div>'
            f'<div class="item-value">R$ {_esc_html(item.get("total",""))}</div>'
            f'</div>'
        )

    # Financeiro
    fin_html = f'<div class="fin-row"><span class="fin-l">Subtotal</span><span class="fin-v">R$ {_fmt_brl(subtotal_v)}</span></div>'
    if desconto_v > 0:
        pct_str = f"{desconto_pct:.2f}".replace(".", ",")
        fin_html += f'<div class="fin-row"><span class="fin-l">Desconto {pct_str}%</span><span class="fin-v">&minus;R$ {_fmt_brl(desconto_v)}</span></div>'
    fin_html += f'<hr class="fin-rule"><div class="fin-total-row"><span class="fin-total-l">Total</span><span class="fin-total-v">R$ {_fmt_brl(total_final_v)}</span></div>'

    # Pagamento e entrega (bottom)
    pgto_detail = ""
    if _h(pix):
        pgto_detail += f"Chave: {_h(pix)}<br>"
    if _h(vencimento):
        pgto_detail += f"Vencimento: {_h(vencimento)}"
    end_entrega  = _h(end_db or obra.get("endereco_entrega", ""))

    # Fornecedor meta
    forn_meta_parts = []
    if _h(cnpj):
        forn_meta_parts.append(f"CNPJ {_h(cnpj)}")
    if forn_local:
        forn_meta_parts.append(_esc_html(forn_local))
    forn_meta = " &nbsp;&middot;&nbsp; ".join(forn_meta_parts)

    ramo_html  = f'<div class="supplier-trade">{_esc_html(ramo)}</div>' if _h(ramo) else ""
    forn_meta_html = f'<div class="supplier-meta">{forn_meta}</div>' if forn_meta else ""

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<title>Pedido de Compra #{pfm_codigo}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&display=swap" rel="stylesheet">
<style>{_PC_CSS}</style>
</head>
<body>
<div class="page">
  <div class="header">
    <div>
      <div class="company-brand">Verschoor Investimentos Imobiliários</div>
      <div class="company-meta">CNPJ {DELTAD['cnpj']} &nbsp;&middot;&nbsp; I.E. {DELTAD['ie']}<br>{_esc_html(DELTAD['end'])}</div>
    </div>
    <div class="doc-meta">
      <div class="doc-tipo">Pedido de Compra</div>
      <div class="doc-number">#{pfm_codigo}</div>
      <div class="doc-date">{data_emissao}</div>
    </div>
  </div>
  <hr class="rule rule-gap">
  <div class="context-block">
    <div><div class="ctx-label">Origem</div><div class="ctx-value">{origem_html}</div></div>
    <div><div class="ctx-label">Entrega</div><div class="ctx-value">{entrega_html}</div></div>
  </div>
  <hr class="rule rule-gap">
  <div class="section-label" style="margin-bottom:6px;">Fornecedor</div>
  <div class="supplier">
    <div class="supplier-name">{_esc_html(fornecedor)}</div>
    {ramo_html}
    {forn_meta_html}
  </div>
  <hr class="rule" style="margin-top:32px;margin-bottom:28px;">
  <div class="section-label">Itens solicitados</div>
  {items_html}
  <div class="financial-outer">
    <div class="financial-inner">{fin_html}</div>
  </div>
  <hr class="rule" style="margin-top:28px;margin-bottom:28px;">
  <div class="bottom">
    <div>
      <div class="bottom-label">Pagamento</div>
      <div class="bottom-main">{_h(condicao) or 'PIX à vista'}</div>
      <div class="bottom-detail">{pgto_detail}</div>
    </div>
    <div>
      <div class="bottom-label">Entrega</div>
      <div class="bottom-main">{_h(data_ent_db) or '—'}</div>
      <div class="bottom-detail">{end_entrega}</div>
    </div>
  </div>
  <div class="footer-tagline">Laura não é uma ferramenta que você usa. É uma memória que você carrega.</div>
</div>
</body>
</html>"""

async def _html_para_pdf(html_str: str) -> bytes:
    from playwright.async_api import async_playwright
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page    = await browser.new_page()
        await page.set_content(html_str, wait_until="networkidle")
        pdf     = await page.pdf(format="A4", print_background=True)
        await browser.close()
        return pdf

def gerar_pfm(doc_id, categoria=None, pfm_codigo_override=None):
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(
            "SELECT ggv, dados_claude, condicao_pgto, data_entrega, endereco_entrega, desconto_rs, caminho FROM documentos WHERE id=?",
            (doc_id,)
        ).fetchone()
    if row is None:
        raise ValueError(f"Documento {doc_id} não encontrado no banco.")
    ggv, dados, condicao, data_entrega_db, endereco, desconto_rs, caminho_original = row

    nome_claude   = _campo(dados, "Fornecedor")
    cnpj_claude   = _campo(dados, "CNPJ/CPF")
    ramo_claude   = _campo(dados, "Ramo de atividade")
    resumo_claude = _campo(dados, "Resumo da compra")
    forn_db       = buscar_fornecedor(nome_claude, cnpj_claude)

    if forn_db:
        fornecedor = forn_db.get("razao_social") or forn_db.get("nome") or nome_claude
        cnpj       = forn_db.get("cnpj") or forn_db.get("cpf") or cnpj_claude
        pix        = forn_db.get("chave_pix") or _campo(dados, "Chave PIX")
        forn_logr  = " ".join(filter(None, [forn_db.get("logradouro"), forn_db.get("numero")]))
        forn_bairro = forn_db.get("bairro") or ""
        # Valida cidade: máx 30 chars, sem '/', sem dígitos — filtra dados errados do import
        _cidade = forn_db.get("cidade") or ""
        _uf     = forn_db.get("uf") or ""
        _cep    = forn_db.get("cep") or ""
        if len(_cidade) > 30 or "/" in _cidade or any(c.isdigit() for c in _cidade):
            _cidade = _uf = _cep = ""
        forn_cidade  = " / ".join(filter(None, [_cidade, _uf, _cep]))
        forn_email   = forn_db.get("email") or ""
        forn_fone    = forn_db.get("whatsapp") or ""
        forn_contato = forn_db.get("contato") or ""
        ramo         = forn_db.get("ramo") or ramo_claude
        # Persiste ramo se ainda não estava cadastrado
        if not forn_db.get("ramo") and ramo_claude and ramo_claude != "A PREENCHER":
            cnpj_key = forn_db.get("cnpj") or forn_db.get("cpf")
            if cnpj_key:
                with sqlite3.connect(DB_PATH) as _con:
                    _con.execute("UPDATE fornecedores SET ramo=? WHERE cnpj=? OR cpf=?",
                                 (ramo_claude, cnpj_key, cnpj_key))
    else:
        fornecedor  = nome_claude
        cnpj        = cnpj_claude
        pix         = _campo(dados, "Chave PIX")
        ramo        = ramo_claude
        forn_logr = forn_bairro = forn_cidade = forn_email = forn_fone = forn_contato = ""
        _criar_fornecedor_auto(nome_claude, cnpj_claude, ramo_claude, doc_id)

    prazo        = _campo(dados, "Prazo de entrega")
    if prazo == "A PREENCHER":
        prazo = _campo(dados, "Data/prazo de entrega")
    data_entrega = data_entrega_db or "A PREENCHER"
    itens        = _itens(dados)
    obs          = _obs(dados)

    subtotal_v, desconto_v, total_final_v = _calcular_totais(dados, desconto_rs)
    valor = f"R$ {_fmt_brl(total_final_v)}" if total_final_v > 0 else _campo(dados, "Valor total")

    if pfm_codigo_override:
        pfm_codigo = pfm_codigo_override
    else:
        pfm_num    = proximo_pfm_numero(ggv)
        pfm_codigo = f"{ggv}-{pfm_num:03d}"
        atualizar(doc_id, pfm_numero=pfm_num, status="pfm_gerado")

    now = datetime.now()
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin   = Cm(2)
        s.right_margin  = Cm(2)
    # Largura útil: 21 - 4 = 17 cm

    # ── CABEÇALHO ────────────────────────────────────────────────────────────
    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.style = "Table Grid"
    _set_col_widths(tbl_h, [11, 6])
    for c in tbl_h.rows[0].cells:
        _cell_bg(c, "D9E2F3")

    c_l = tbl_h.rows[0].cells[0]
    r = c_l.paragraphs[0].add_run("DeltaD Engenharia")
    r.bold = True; r.font.size = Pt(13)
    p2 = c_l.add_paragraph("Pedido de Compra")
    p2.runs[0].font.size = Pt(9)

    c_r = tbl_h.rows[0].cells[1]
    c_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = c_r.paragraphs[0].add_run(f"Nº: {pfm_codigo}")
    r2.bold = True; r2.font.size = Pt(12)
    p3 = c_r.add_paragraph(_data_extenso(now))
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.runs[0].font.size = Pt(8)

    # ── FORNECEDOR ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    tbl_f = doc.add_table(rows=0, cols=2)
    tbl_f.style = "Table Grid"
    _set_col_widths(tbl_f, [5, 12])
    _secao_row(tbl_f, "FORNECEDOR", 2)
    _kv_row(tbl_f, "RAZÃO SOCIAL / NOME", fornecedor)
    _kv_row(tbl_f, "CNPJ/CPF", cnpj)
    _kv_row(tbl_f, "I.E.", "ISENTO")
    if forn_logr:
        _kv_row(tbl_f, "LOGRADOURO / Nº", forn_logr)
    if forn_bairro:
        _kv_row(tbl_f, "BAIRRO", forn_bairro)
    if forn_cidade:
        _kv_row(tbl_f, "CIDADE / UF / CEP", forn_cidade)
    if forn_contato:
        _kv_row(tbl_f, "CONTATO", forn_contato)
    if forn_email:
        _kv_row(tbl_f, "E-MAIL", forn_email)
    if forn_fone:
        _kv_row(tbl_f, "WHATSAPP", forn_fone)
    _kv_row(tbl_f, "CHAVE PIX", pix)

    # ── EMPREENDIMENTO ───────────────────────────────────────────────────────
    doc.add_paragraph()
    tbl_e = doc.add_table(rows=0, cols=2)
    tbl_e.style = "Table Grid"
    _set_col_widths(tbl_e, [5, 12])
    _secao_row(tbl_e, "EMPREENDIMENTO", 2)
    _kv_row(tbl_e, ggv, GGV_DESC.get(ggv, ggv))

    # ── MATERIAIS ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    tbl_m = doc.add_table(rows=0, cols=6)
    tbl_m.style = "Table Grid"
    _set_col_widths(tbl_m, [1.2, 7.3, 1.5, 1.5, 2.75, 2.75])
    _secao_row(tbl_m, "MATERIAIS", 6)

    # Cabeçalho de colunas
    row_hdr = tbl_m.add_row()
    for i, h in enumerate(["ID", "DESCRIÇÃO", "UND", "QTDE", "R$ UNIT", "R$ TOTAL"]):
        c = row_hdr.cells[i]
        _cell_bg(c, "D9D9D9")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(8)

    # Itens
    lista_itens = itens if itens else [dados[:200]]
    for idx, item in enumerate(lista_itens, 1):
        ri = tbl_m.add_row()
        ri.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        ri.cells[0].paragraphs[0].add_run(f"{idx:02d}").font.size = Pt(9)
        if isinstance(item, dict):
            ri.cells[1].paragraphs[0].add_run(item["desc"]).font.size = Pt(9)
            for j, key in [(2, "und"), (3, "qtde"), (4, "unit"), (5, "total")]:
                ri.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ri.cells[j].paragraphs[0].add_run(str(item[key])).font.size = Pt(9)
        else:
            ri.cells[1].paragraphs[0].add_run(str(item).lstrip("- ")).font.size = Pt(9)
            for j in [2, 3, 4, 5]:
                ri.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                ri.cells[j].paragraphs[0].add_run("—").font.size = Pt(9)

    # Totais — subtotal + desconto (se houver) + total final
    def _total_row(tbl, label, valor_str, bg="D9D9D9"):
        r = tbl.add_row()
        c_label = r.cells[0]
        for i in range(1, 5):
            c_label.merge(r.cells[i])
        _cell_bg(c_label, bg)
        p = c_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        run.bold = True; run.font.size = Pt(9)
        c_v = r.cells[5]
        _cell_bg(c_v, bg)
        p_v = c_v.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_v = p_v.add_run(valor_str)
        run_v.bold = True; run_v.font.size = Pt(9)

    if desconto_v > 0 and subtotal_v > 0:
        pct = desconto_v / subtotal_v * 100
        _total_row(tbl_m, "SUBTOTAL:", f"R$ {_fmt_brl(subtotal_v)}", "F2F2F2")
        _total_row(tbl_m, f"DESCONTO ({pct:.2f}%):", f"-R$ {_fmt_brl(desconto_v)}", "F2F2F2")

    _total_row(tbl_m, "TOTAL DO PEDIDO:", valor)

    # ── PARTE INFERIOR: PRAZO | DADOS ────────────────────────────────────────
    doc.add_paragraph()
    tbl_b = doc.add_table(rows=1, cols=2)
    tbl_b.style = "Table Grid"
    _set_col_widths(tbl_b, [8.5, 8.5])

    sem_entrega = bool(categoria) and categoria.value in CATEGORIAS_SEM_NFE_OBRIGATORIA

    # — Esquerda: Prazo e Condições —
    c_prazo = tbl_b.rows[0].cells[0]
    p = c_prazo.paragraphs[0]
    titulo_prazo = "CONDIÇÕES DE PAGAMENTO" if sem_entrega else "PRAZO PARA ENTREGA E CONDIÇÕES DE PAGAMENTO"
    r = p.add_run(titulo_prazo)
    r.bold = True; r.font.size = Pt(8)
    _cell_bg(c_prazo, "D9D9D9")

    def _kv_p(cell, label, val):
        p = cell.add_paragraph()
        rl = p.add_run(f"{label}: ")
        rl.bold = True; rl.font.size = Pt(8)
        p.add_run(str(val or "A PREENCHER")).font.size = Pt(9)

    _kv_p(c_prazo, "CONDIÇÕES DE PAGAMENTO", condicao)
    _kv_p(c_prazo, "CHAVE PIX", pix)
    if not sem_entrega:
        _kv_p(c_prazo, "DATA DE ENTREGA", data_entrega)
    obs_partes = []
    prazo_texto = prazo if not sem_entrega and prazo and prazo not in ("A PREENCHER",) and prazo != data_entrega else None
    if prazo_texto:
        obs_partes.append(prazo_texto)
    if obs and obs != prazo_texto:
        obs_partes.append(obs)
    if obs_partes:
        _kv_p(c_prazo, "OBSERVAÇÃO", " | ".join(obs_partes))

    if not sem_entrega:
        # Nota de foto
        p_foto = c_prazo.add_paragraph()
        p_foto.add_run(
            "FAVOR TIRAR FOTOS DO MATERIAL DESCARREGADO E ENVIAR POR WHATSAPP PARA DENNIS – (42) 99127-1255"
        ).font.size = Pt(7)

    # — Direita: Dados para Fatura (+ Entrega, quando aplicável) —
    c_dados = tbl_b.rows[0].cells[1]
    p_fatura = c_dados.paragraphs[0]
    r_f = p_fatura.add_run("DADOS PARA FATURA")
    r_f.bold = True; r_f.font.size = Pt(8)
    _cell_bg(c_dados, "EBF3F0")

    def _linha(cell, txt, bold=False, size=9):
        p = cell.add_paragraph()
        r = p.add_run(txt)
        r.bold = bold; r.font.size = Pt(size)

    _linha(c_dados, DELTAD["nome"], bold=True, size=9)
    _linha(c_dados, f"CNPJ: {DELTAD['cnpj']}", size=9)
    _linha(c_dados, DELTAD["end"], size=8)
    _linha(c_dados, DELTAD["email"], size=8)
    if not sem_entrega:
        _linha(c_dados, "")
        p_ent = c_dados.add_paragraph()
        p_ent.add_run("DADOS PARA ENTREGA").bold = True
        p_ent.runs[0].font.size = Pt(8)
        _linha(c_dados, endereco or "A PREENCHER", size=9)

    pasta = _pasta_pfm(ggv)
    pasta.mkdir(parents=True, exist_ok=True)
    prefixo   = "TESTE-" if TEST_MODE else ""
    nome_base = _nome_base_pfm(pfm_codigo, fornecedor, resumo_claude, prefixo)
    caminho   = pasta / f"{nome_base}.docx"
    doc.save(caminho)

    if pfm_codigo_override:
        lanc_status, ja_existia = "a_pagar", True
    else:
        lanc_status, ja_existia = registrar_lancamento(
            doc_id, pfm_codigo, ggv, fornecedor, total_final_v, data_entrega_db, categoria
        )
        atualizar(doc_id, caminho_pfm=str(caminho))
        # Arquiva o orçamento original em "00 Orçamentos" — só na geração original, não em revisões
        if caminho_original and Path(caminho_original).exists():
            ext_original = Path(caminho_original).suffix
            destino = _pasta_orcamentos(ggv) / f"{nome_base}{ext_original}"
            try:
                shutil.copy2(caminho_original, destino)
            except OSError:
                pass
    return caminho, pfm_codigo, fornecedor, total_final_v, lanc_status, ja_existia

# ── Comprovante PIX ────────────────────────────────────────────────────────

def parse_comprovante(dados_claude: str) -> dict:
    """Extrai campos estruturados do texto retornado pelo Claude para comprovante_pix."""
    valor_str = _campo(dados_claude, "Valor")
    try:
        valor_v = _parse_brl(re.sub(r"[^\d,.]", "", valor_str)) if valor_str != "A PREENCHER" else 0.0
    except Exception:
        valor_v = 0.0
    return {
        "valor_v":      valor_v,
        "valor_fmt":    _fmt_brl(valor_v) if valor_v > 0 else valor_str,
        "data":         _campo(dados_claude, "Data do pagamento"),
        "favorecido":   _campo(dados_claude, "Favorecido"),
        "cnpj":         _campo(dados_claude, "CNPJ/CPF do favorecido"),
        "chave_pix":    _campo(dados_claude, "Chave PIX"),
        "instituicao":  _campo(dados_claude, "Instituição financeira"),
        "id_transacao": _campo(dados_claude, "ID da transação"),
        "obs":          _campo(dados_claude, "Identificador / Observação"),
    }

def buscar_candidatos_pix(valor_v: float, favorecido: str, cnpj: str) -> list:
    """Pontua lançamentos A PAGAR e retorna os 3 melhores candidatos."""
    cnpj_digits = re.sub(r"\D", "", cnpj) if cnpj and cnpj != "A PREENCHER" else ""
    fav_token   = favorecido.strip().upper().split()[0] if favorecido and favorecido != "A PREENCHER" else ""

    # Tenta validar o nome via CNPJ na tabela de fornecedores
    nome_canonico_token = ""
    if cnpj_digits:
        with sqlite3.connect(DB_PATH) as con:
            row = con.execute(
                "SELECT nome FROM fornecedores "
                "WHERE REPLACE(REPLACE(REPLACE(cnpj,'.',''),'/',''),'-','') = ? LIMIT 1",
                (cnpj_digits,)
            ).fetchone()
        if row:
            nome_canonico_token = row[0].strip().upper().split()[0]

    with sqlite3.connect(DB_PATH) as con:
        rows = con.execute(
            "SELECT pfm_codigo, ggv, fornecedor, valor FROM lancamentos WHERE status='a_pagar'"
        ).fetchall()

    candidatos = []
    for pfm_codigo, ggv, fornecedor, valor_lanc in rows:
        score     = 0
        forn_token = (fornecedor or "").strip().upper().split()[0] if fornecedor else ""

        # Valor
        if valor_v > 0 and valor_lanc:
            if abs(valor_v - valor_lanc) <= 0.01:
                score += 3
            elif valor_lanc > 0 and abs(valor_v - valor_lanc) / valor_lanc <= 0.10:
                score += 1

        # Nome — CNPJ validado tem peso maior que coincidência direta
        if nome_canonico_token and forn_token == nome_canonico_token:
            score += 3
        elif fav_token and forn_token == fav_token:
            score += 2

        if score > 0:
            candidatos.append({
                "pfm_codigo": pfm_codigo,
                "ggv":        ggv,
                "fornecedor": fornecedor,
                "valor_lanc": valor_lanc,
                "score":      score,
            })

    candidatos.sort(key=lambda c: c["score"], reverse=True)
    return candidatos[:3]

def mostrar_comprovante_candidatos(dados: dict, candidatos: list) -> str:
    """Formata a mensagem de comprovante + candidatos para o Telegram. Sem IO."""
    def _conf(score):
        if score >= 5: return "exato"
        if score >= 3: return "próximo"
        return ""

    linhas = ["Pagamento identificado.\n"]
    if dados["favorecido"] != "A PREENCHER":
        linhas.append(f"{dados['favorecido']} — R$ {dados['valor_fmt']}")
    else:
        linhas.append(f"R$ {dados['valor_fmt']}")
    if dados["data"]        != "A PREENCHER": linhas.append(dados["data"])
    if dados["instituicao"] != "A PREENCHER": linhas.append(dados["instituicao"])
    if dados["obs"]         != "A PREENCHER": linhas.append(f"Ref: {dados['obs']}")

    linhas.append("")

    if not candidatos:
        linhas.append("Nenhum pedido em aberto corresponde a este pagamento.")
        return "\n".join(linhas)

    linhas.append("Qual pedido este pagamento quita?\n")
    for c in candidatos:
        valor_fmt = f"R$ {_fmt_brl(c['valor_lanc'])}" if c["valor_lanc"] else "—"
        conf = _conf(c["score"])
        linha = f"🟡 #{c['pfm_codigo']} · {c['fornecedor']} · {valor_fmt}"
        if conf:
            linha += f" · {conf}"
        linhas.append(linha)

    return "\n".join(linhas)

# ── UI helpers ─────────────────────────────────────────────────────────────

def parse_resposta(texto):
    tipo, ggv = "nao_relacionado", "nao_identificado"
    corpo = []
    for linha in texto.strip().splitlines():
        if linha.startswith("TIPO:"):
            tipo = linha.split(":", 1)[1].strip().strip("[]").split("|")[0].strip()
        elif linha.startswith("GGV:"):
            ggv = linha.split(":", 1)[1].strip().strip("[]").split("|")[0].strip()
        else:
            corpo.append(linha)
    return tipo, ggv, "\n".join(corpo).strip()

def mostrar_cockpit_obra(obra):
    codigo    = obra.get("codigo", "")
    desc      = obra.get("descricao", "") or ""
    enc_nome  = obra.get("encarregado_nome", "") or ""
    enc_fone  = obra.get("encarregado_fone", "") or ""
    resp_nome = obra.get("responsavel_nome", "") or ""
    resp_fone = obra.get("responsavel_fone", "") or ""
    end       = obra.get("endereco_entrega", "") or ""

    # Header: "GGV03 — Condomínio residencial" + detalhe em linha separada
    if "," in desc:
        titulo, detalhe = desc.split(",", 1)
        titulo  = titulo.replace(codigo, "").strip()
        detalhe = detalhe.strip().rstrip(".")
    else:
        titulo  = desc.replace(codigo, "").strip() or "Sem descrição"
        detalhe = ""
    cabecalho = f"{codigo} — {titulo}"
    if detalhe:
        cabecalho += f"\n{detalhe}"

    SEP = "──────────────────────────────"

    # Placeholder financeiro — Fiada 5b-1
    financeiro = "⚪ Nenhum lançamento registrado"

    # Contatos com separador · consistente
    enc_txt  = f"Encarregado · {enc_nome} · {enc_fone}".strip(" ·") if enc_nome else "Encarregado · —"
    resp_txt = f"Responsável · {resp_nome} · {resp_fone}".strip(" ·") if resp_nome else "Responsável · —"

    # Endereço: remove CEP e troca " - " por " · "
    end_curto = re.sub(r"\s+CEP[\s\d\.\-]+$", "", end).replace(" - ", " · ") if end else ""

    contatos = f"{enc_txt}\n{resp_txt}"
    if end_curto:
        contatos += f"\nEntrega · {end_curto}"

    return f"{cabecalho}\n\n{SEP}\n{financeiro}\n{SEP}\n{contatos}"

_SAUDACOES_RE = re.compile(
    r"^\s*(oi|olá|ola|bom\s*dia|boa\s*tarde|boa\s*noite|hey|e\s*a[íi]|hello|hi|tudo\s*bem|tudo\s*bom)\W*$",
    re.IGNORECASE
)
_OBRAS_RE = re.compile(r"^\s*obras?\s*$", re.IGNORECASE)

def mostrar_boas_vindas():
    return (
        "Por onde quer começar?\n\n"
        "Obras — todas as obras e pedidos\n"
        "Ajuda — pedidos de compra, pagamentos e consultas"
    )

def teclado_boas_vindas():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📋 Obras",  callback_data="menu_obras")],
        [InlineKeyboardButton("❓ Ajuda",  callback_data="menu_ajuda")],
        [InlineKeyboardButton("✖ Fechar", callback_data="obras_fechar")],
    ])

def mostrar_ajuda():
    return (
        "No que posso ajudar?\n\n"
        "<b>Cadastrar pedido de compra</b>\n"
        "Envie a foto ou arquivo do orçamento.\n\n"
        "<b>Confirmar pagamento</b>\n"
        "Envie o comprovante PIX.\n\n"
        "<b>Incluir nota fiscal</b>\n"
        "Envie o PDF ou foto da NF-e.\n\n"
        "<b>Registrar entrega</b>\n"
        "Envie a foto ou use /entrega. Também disponível no botão 📦 Entregue dentro do pedido.\n\n"
        "<b>Consultas diretas</b>\n"
        "Digite o código da obra (GGV03) ou do pedido (GGV03-009)."
    )

def teclado_ajuda():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("← Voltar",  callback_data="menu_inicio")],
        [InlineKeyboardButton("✖ Fechar", callback_data="obras_fechar")],
    ])

def _listar_obras():
    with sqlite3.connect(DB_PATH) as con:
        return con.execute(
            "SELECT codigo, descricao FROM obras WHERE ativa=1 ORDER BY codigo"
        ).fetchall()

def mostrar_lista_obras(obras):
    if not obras:
        return "Nenhuma obra cadastrada. Use ➕ Nova obra para começar."
    linhas = ["Qual obra?", ""]
    for codigo, desc in obras:
        desc = desc or ""
        titulo = desc.split(",")[0].replace(codigo, "").strip() if desc else "Sem descrição"
        linhas.append(f"{codigo} — {titulo}")
    return "\n".join(linhas)

def teclado_lista_obras(obras):
    botoes = []
    row = []
    for codigo, _ in obras:
        row.append(InlineKeyboardButton(codigo, callback_data=f"obra_ver:{codigo}"))
        if len(row) == 2:
            botoes.append(row)
            row = []
    if row:
        botoes.append(row)
    botoes.append([InlineKeyboardButton("➕ Nova obra",  callback_data="menu_nova_obra")])
    botoes.append([InlineKeyboardButton("← Voltar",      callback_data="menu_inicio")])
    botoes.append([InlineKeyboardButton("✖ Fechar",     callback_data="obras_fechar")])
    return InlineKeyboardMarkup(botoes)

def _pedidos_obra(ggv):
    with sqlite3.connect(DB_PATH) as con:
        return con.execute(
            "SELECT pfm_codigo, status, fornecedor, valor FROM lancamentos WHERE ggv=? ORDER BY pfm_codigo",
            (ggv,)
        ).fetchall()

def teclado_obra(codigo, pedidos=None):
    botoes = []
    if pedidos:
        botoes.append([InlineKeyboardButton("📋 Pedidos",   callback_data=f"obra_pedidos:{codigo}")])
    botoes.append([InlineKeyboardButton("✏️ Editar obra",   callback_data=f"obra_editar:{codigo}")])
    botoes.append([InlineKeyboardButton("◀️ Obras",         callback_data="menu_obras")])
    botoes.append([InlineKeyboardButton("✖ Fechar",         callback_data=f"obra_fechar:{codigo}")])
    return InlineKeyboardMarkup(botoes)

def mostrar_lista_pedidos(codigo, pedidos):
    _ST = {"a_pagar": "🟡", "pago": "🟢", "pendente_revisao": "🔴", "substituido": "⚫"}
    if not pedidos:
        return f"Nenhum pedido em {codigo}. Envie um orçamento para começar."
    linhas = [f"Qual pedido? · {codigo}\n"]
    for pfm_codigo, status, fornecedor, valor in pedidos:
        emoji    = _ST.get(status, "⚪")
        forn_cur = (fornecedor or "—")[:20]
        val_str  = f"R$ {_fmt_brl(valor)}" if valor else "—"
        linhas.append(f"{emoji} {pfm_codigo}  {forn_cur} · {val_str}")
    return "\n".join(linhas)

def teclado_lista_pedidos(codigo, pedidos):
    _ST = {"a_pagar": "🟡", "pago": "🟢", "pendente_revisao": "🔴", "substituido": "⚫"}
    botoes = []
    row = []
    for pfm_codigo, status, *_ in pedidos:
        emoji = _ST.get(status, "⚪")
        row.append(InlineKeyboardButton(f"{emoji} {pfm_codigo}", callback_data=f"pedido_abrir:{pfm_codigo}"))
        if len(row) == 2:
            botoes.append(row)
            row = []
    if row:
        botoes.append(row)
    botoes.append([InlineKeyboardButton("◀️ Voltar à obra", callback_data=f"obra_ver:{codigo}")])
    return InlineKeyboardMarkup(botoes)

def teclado_obra_campos(codigo):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("Descrição",             callback_data=f"obra_campo:{codigo}:descricao")],
        [InlineKeyboardButton("Endereço de entrega",   callback_data=f"obra_campo:{codigo}:endereco_entrega")],
        [InlineKeyboardButton("Encarregado (nome)",    callback_data=f"obra_campo:{codigo}:encarregado_nome")],
        [InlineKeyboardButton("Encarregado (fone)",    callback_data=f"obra_campo:{codigo}:encarregado_fone")],
        [InlineKeyboardButton("Responsável (nome)",    callback_data=f"obra_campo:{codigo}:responsavel_nome")],
        [InlineKeyboardButton("Responsável (fone)",    callback_data=f"obra_campo:{codigo}:responsavel_fone")],
        [InlineKeyboardButton("◀️ Voltar",             callback_data=f"obra_ver:{codigo}")],
    ])

def teclado_orcamento(doc_id, tipo, ggv):
    if ggv == "nao_identificado":
        return InlineKeyboardMarkup([
            [InlineKeyboardButton("📍 Definir obra",   callback_data=f"sel_ggv:{doc_id}:{tipo}:{ggv}")],
            [InlineKeyboardButton("✏️ Corrigir dados", callback_data=f"sel_edit:{doc_id}:{tipo}:{ggv}")],
            [InlineKeyboardButton("Cancelar",          callback_data=f"cancelar:{doc_id}")],
        ])
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✅ Gerar Pedido de Compra", callback_data=f"pfm:{doc_id}:{ggv}")],
        [InlineKeyboardButton("✏️ Corrigir dados",         callback_data=f"sel_edit:{doc_id}:{tipo}:{ggv}")],
        [InlineKeyboardButton("Cancelar",                  callback_data=f"cancelar:{doc_id}")],
    ])

def teclado_candidatos_pix(doc_id_comp: int, candidatos: list):
    botoes = []
    for c in candidatos:
        botoes.append([InlineKeyboardButton(
            f"Pedido #{c['pfm_codigo']}",
            callback_data=f"pix_confirmar:{doc_id_comp}:{c['pfm_codigo']}"
        )])
    botoes.append([InlineKeyboardButton("Nenhum destes", callback_data="pix_cancelar")])
    return InlineKeyboardMarkup(botoes)

def teclado_tipo_inicial(doc_id):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📋 Orçamento / Cotação",  callback_data=f"sel_tipo_inicial:{doc_id}:orcamento")],
        [InlineKeyboardButton("💰 Comprovante PIX",       callback_data=f"sel_tipo_inicial:{doc_id}:comprovante_pix")],
        [InlineKeyboardButton("🧾 Nota Fiscal",           callback_data=f"sel_tipo_inicial:{doc_id}:nota_fiscal")],
        [InlineKeyboardButton("📦 Foto de entrega",       callback_data=f"sel_tipo_inicial:{doc_id}:foto_entrega")],
        [InlineKeyboardButton("🏦 Extrato Mercado Pago", callback_data=f"sel_tipo_inicial:{doc_id}:extrato_mp")],
        [InlineKeyboardButton("Não é da obra",            callback_data=f"sel_tipo_inicial:{doc_id}:nao_relacionado")],
    ])

def teclado_condicao(doc_id, tipo, ggv):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("💰 PIX à vista",                   callback_data=f"pgto:{doc_id}:{ggv}:pix_avista")],
        [InlineKeyboardButton("💰 PIX 50% entrada + 50% entrega", callback_data=f"pgto:{doc_id}:{ggv}:pix_50_50")],
        [InlineKeyboardButton("✏️ Outro (digitar)",                callback_data=f"pgto:{doc_id}:{ggv}:outro")],
        [InlineKeyboardButton("◀️ Voltar",                        callback_data=f"voltar_edit:{doc_id}:{tipo}:{ggv}")],
    ])

def teclado_endereco(doc_id, tipo, ggv):
    chave_obra = f"obra_{ggv}"
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(f"🏗 Obra ({ggv})",    callback_data=f"end:{doc_id}:{ggv}:{chave_obra}")],
        [InlineKeyboardButton("🏠 Casa",             callback_data=f"end:{doc_id}:{ggv}:casa")],
        [InlineKeyboardButton("🏢 Escritório",       callback_data=f"end:{doc_id}:{ggv}:escritorio")],
        [InlineKeyboardButton("🌳 Chácara",          callback_data=f"end:{doc_id}:{ggv}:chacara")],
        [InlineKeyboardButton("✏️ Outro (digitar)",  callback_data=f"end:{doc_id}:{ggv}:outro")],
        [InlineKeyboardButton("◀️ Voltar",           callback_data=f"voltar_edit:{doc_id}:{tipo}:{ggv}")],
    ])

def _fmt_data_curta(dt_str):
    """'2026-07-03 10:30:00' → '03/07'"""
    try:
        return f"{dt_str[8:10]}/{dt_str[5:7]}"
    except Exception:
        return dt_str[:10] if dt_str else "—"

def buscar_pedido(pfm_codigo: str) -> Optional[Pedido]:
    """Consulta o banco e retorna um Pedido com dados brutos e cálculos financeiros, ou None."""
    m = re.match(r"^(GGV\d{2,3})-(\d{3})(?:-R\d+)?$", pfm_codigo.upper())
    if not m:
        return None
    ggv, pfm_num = m.group(1), int(m.group(2))
    with sqlite3.connect(DB_PATH) as con:
        doc = con.execute(
            "SELECT id, ggv, dados_claude, condicao_pgto, data_entrega, desconto_rs, caminho, criado_em "
            "FROM documentos WHERE ggv=? AND pfm_numero=?",
            (ggv, pfm_num)
        ).fetchone()
        if not doc:
            return None
        doc_id, ggv_db, dados, condicao_pgto, data_entrega, desconto_rs, caminho, doc_criado = doc
        lanc = con.execute(
            "SELECT fornecedor, valor, data_prevista_entrega, vencimento_pagamento, status, criado_em, "
            "data_pagamento, doc_id_nfe, doc_id_comprovante, identificador_comprovante, "
            "obs_entrega, entregue_em, categoria "
            "FROM lancamentos WHERE pfm_codigo=?",
            (pfm_codigo,)
        ).fetchone()
        qtd_fotos = con.execute(
            "SELECT COUNT(*) FROM entrega_fotos WHERE pfm_codigo=?", (pfm_codigo,)
        ).fetchone()[0]

    forn_lanc = data_prev_ent = venc = status_raw = lanc_criado = data_pgto = None
    doc_id_nfe = doc_id_comp = ident_comp = obs_ent = entregue_em = categoria_lanc = None
    if lanc:
        forn_lanc, _, data_prev_ent, venc, status_raw, lanc_criado, data_pgto, doc_id_nfe, doc_id_comp, ident_comp, obs_ent, entregue_em, categoria_lanc = lanc

    try:
        status = StatusPedido(status_raw) if status_raw else StatusPedido.SEM_LANCAMENTO
    except ValueError:
        status = StatusPedido.SEM_LANCAMENTO

    subtotal_v, desconto_v, total_v = _calcular_totais(dados, desconto_rs)

    return Pedido(
        codigo             = pfm_codigo,
        doc_id             = doc_id,
        ggv                = ggv_db,
        status             = status,
        fornecedor         = forn_lanc or _campo(dados, "Fornecedor") or "—",
        cnpj               = _campo(dados, "CNPJ/CPF") or "—",
        valor_orcamento    = subtotal_v,
        desconto           = desconto_v,
        valor_negociado    = total_v,
        condicao_pagamento = condicao_pgto or "—",
        vencimento         = venc or "—",
        entrega_prevista   = data_prev_ent or data_entrega or "—",
        doc_criado_em      = doc_criado,
        lanc_criado_em     = lanc_criado,
        data_pagamento     = data_pgto,
        doc_id_nfe                = doc_id_nfe,
        doc_id_comprovante        = doc_id_comp,
        identificador_comprovante = ident_comp,
        qtd_fotos_entrega         = qtd_fotos,
        obs_entrega               = obs_ent,
        entregue_em               = entregue_em,
        categoria                 = categoria_lanc,
        caminho_orcamento         = caminho,
    )

def preparar_visualizacao_pedido(pedido: Pedido) -> Pedido:
    """Verifica existência de arquivos em disco e constrói o histórico. Retorna o Pedido enriquecido."""
    if pedido.caminho_orcamento and not Path(pedido.caminho_orcamento).exists():
        pedido.caminho_orcamento = None

    with sqlite3.connect(DB_PATH) as con:
        row = con.execute("SELECT caminho_pfm FROM documentos WHERE id=?", (pedido.doc_id,)).fetchone()
    caminho_pfm = row[0] if row else None
    pedido.caminho_docx = caminho_pfm if caminho_pfm and Path(caminho_pfm).exists() else None

    if pedido.doc_id_nfe:
        with sqlite3.connect(DB_PATH) as con:
            row = con.execute("SELECT dados_claude FROM documentos WHERE id=?", (pedido.doc_id_nfe,)).fetchone()
        if row:
            numero = _campo(row[0], "Número da NF")
            pedido.nfe_numero = numero if numero != "A PREENCHER" else None
            data_raw = _campo(row[0], "Data de emissão")
            pedido.nfe_data = data_raw[:5] if data_raw != "A PREENCHER" else None

    historico = []
    if pedido.lanc_criado_em:
        historico.append((_fmt_data_curta(pedido.lanc_criado_em), "Pedido criado"))
    if pedido.data_pagamento:
        dt = pedido.data_pagamento
        if len(dt) >= 5 and dt[2:3] == "/":
            data_fmt = dt[:5]
        else:
            data_fmt = _fmt_data_curta(dt)
        pago_label = "Pago"
        if pedido.identificador_comprovante:
            cod = pedido.identificador_comprovante[:12]
            pago_label = f"Pago · {cod}"
        historico.append((data_fmt, pago_label))
    if pedido.doc_id_nfe:
        nfe_label = f"NF-e {pedido.nfe_numero}" if pedido.nfe_numero else "NF-e"
        historico.append((pedido.nfe_data or "", nfe_label))
    if pedido.obs_entrega:
        ent_data = _fmt_data_curta(pedido.entregue_em) if pedido.entregue_em else ""
        ent_label = "Entregue" if pedido.obs_entrega == "Entrega completa" else f"Entregue — {pedido.obs_entrega}"
        historico.append((ent_data, ent_label))
    pedido.historico = historico

    return pedido

# Categorias cujo fechamento fiscal é a própria fatura (CREA, ONR, prefeitura, Copel, Sanepar
# não emitem NF-e separada) — não exibir "NF-e pendente" nem exigir vínculo de NF-e
CATEGORIAS_SEM_NFE_OBRIGATORIA = {"taxa", "imposto", "servicos"}

def _status_pago_label(pedido: "Pedido") -> str:
    if pedido.nfe_numero:
        return f"Pago · NF-e {pedido.nfe_numero}"
    if pedido.categoria in CATEGORIAS_SEM_NFE_OBRIGATORIA:
        return "Pago"
    if not pedido.doc_id_nfe:
        return "Pago · NF-e pendente"
    return "Pago · NF-e"

def mostrar_pedido(pedido: Pedido) -> str:
    """Formata o Pedido como mensagem Telegram. Sem IO — apenas formatação."""
    _STATUS_EMOJI = {
        StatusPedido.A_PAGAR:          "🟡",
        StatusPedido.PAGO:             "🟢",
        StatusPedido.PENDENTE_REVISAO: "🔴",
        StatusPedido.SUBSTITUIDO:      "⚫",
        StatusPedido.SEM_LANCAMENTO:   "⚪",
    }
    _STATUS_SHORT = {
        StatusPedido.A_PAGAR:          "Aguardando pagamento",
        StatusPedido.PAGO:             _status_pago_label(pedido),
        StatusPedido.PENDENTE_REVISAO: "Requer atenção",
        StatusPedido.SUBSTITUIDO:      "Substituído",
        StatusPedido.SEM_LANCAMENTO:   "Sem registro financeiro",
    }
    SEP = "\n──────────────────────────────\n"

    emoji  = _STATUS_EMOJI.get(pedido.status, "")
    status = _STATUS_SHORT.get(pedido.status, str(pedido.status))
    cabecalho = f"{emoji} #{pedido.codigo} — {status}\n\n{pedido.fornecedor}"

    linhas_fin = []
    if pedido.valor_negociado > 0:
        valor_str = f"R$ {_fmt_brl(pedido.valor_negociado)}"
        if pedido.desconto > 0:
            valor_str += f"  (desc. R$ {_fmt_brl(pedido.desconto)})"
        linhas_fin.append(valor_str)
    cond = pedido.condicao_pagamento if pedido.condicao_pagamento not in ("—", None) else None
    ent  = pedido.entrega_prevista   if pedido.entrega_prevista   not in ("—", None) else None
    if cond and ent:
        linhas_fin.append(f"{cond} · entrega {ent}")
    elif cond:
        linhas_fin.append(cond)
    elif ent:
        linhas_fin.append(f"Entrega: {ent}")
    if pedido.vencimento and pedido.vencimento not in ("—", None):
        linhas_fin.append(f"Vencimento: {pedido.vencimento}")
    financeiro = "\n".join(linhas_fin) if linhas_fin else "Dados financeiros não disponíveis"

    arq = []
    if pedido.caminho_orcamento:
        arq.append("📎 Orçamento original")
    if pedido.caminho_docx:
        arq.append("📄 Pedido de Compra")
    if pedido.doc_id_comprovante:
        arq.append("💰 Comprov. pagamento")
    if pedido.doc_id_nfe:
        nfe_label = f"🧾 NF-e {pedido.nfe_numero}" if pedido.nfe_numero else "🧾 NF-e"
        arq.append(nfe_label)
    if pedido.qtd_fotos_entrega:
        arq.append(f"📦 {_rotulo_qtd_arquivos(pedido.qtd_fotos_entrega)} da entrega")
    arquivos = "\n".join(arq) if arq else "Nenhum arquivo disponível"

    hist = []
    for data, evento in pedido.historico:
        hist.append(f"{data}  {evento}".strip())
    historico = "\n".join(hist) if hist else "—"

    return SEP.join([cabecalho, financeiro, arquivos, historico])

def teclado_pedido(doc_id, pfm_codigo, doc_id_nfe=None, doc_id_comprovante=None,
                   qtd_fotos_entrega=0, obs_entrega=None):
    ggv = pfm_codigo.rsplit("-", 1)[0]
    botoes = [
        [InlineKeyboardButton("Revisar",      callback_data=f"pfm_revisar:{doc_id}:{pfm_codigo}")],
        [InlineKeyboardButton("📄 PDF",       callback_data=f"pfm_ver:{doc_id}:{pfm_codigo}")],
        [InlineKeyboardButton("📎 Orçamento", callback_data=f"pfm_orc:{doc_id}:{pfm_codigo}")],
    ]
    if doc_id_comprovante:
        botoes.append([InlineKeyboardButton("💰 Comprovante", callback_data=f"pfm_comp:{doc_id_comprovante}:{pfm_codigo}")])
    if doc_id_nfe:
        botoes.append([InlineKeyboardButton("🧾 NF-e", callback_data=f"pfm_nfe:{doc_id_nfe}:{pfm_codigo}")])
    if obs_entrega:
        if qtd_fotos_entrega:
            botoes.append([InlineKeyboardButton(
                f"📦 Ver {_rotulo_qtd_arquivos(qtd_fotos_entrega)} da entrega",
                callback_data=f"entrega_ver_fotos:{pfm_codigo}"
            )])
        botoes.append([InlineKeyboardButton("✏️ Editar entrega", callback_data=f"entrega_editar:{pfm_codigo}")])
    else:
        botoes.append([InlineKeyboardButton("📦 Entregue", callback_data=f"pfm_entregue:{doc_id}:{pfm_codigo}")])
    botoes += [
        [InlineKeyboardButton("◀️ Pedidos",   callback_data=f"obra_pedidos:{ggv}")],
        [InlineKeyboardButton("✖ Fechar",     callback_data=f"pfm_fechar:{doc_id}")],
    ]
    return InlineKeyboardMarkup(botoes)

def buscar_pedidos_sem_entrega():
    with sqlite3.connect(DB_PATH) as con:
        return con.execute(
            "SELECT pfm_codigo, fornecedor, valor, status FROM lancamentos "
            "WHERE obs_entrega IS NULL "
            "ORDER BY pfm_codigo DESC"
        ).fetchall()

def _teclado_pedidos_entrega(pedidos):
    botoes = []
    for pfm_codigo, forn, valor, status in pedidos:
        emoji = "🟡" if status == "a_pagar" else "🟢"
        forn_curto = (forn or "")[:22]
        botoes.append([InlineKeyboardButton(
            f"{emoji} #{pfm_codigo} — {forn_curto}",
            callback_data=f"entrega_sel:{pfm_codigo}"
        )])
    botoes.append([InlineKeyboardButton("✖ Cancelar", callback_data="entrega_cancelar")])
    return InlineKeyboardMarkup(botoes)

def teclado_obs_entrega():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✅ Entrega completa",      callback_data="entrega_obs:completa")],
        [InlineKeyboardButton("📦 Entrega parcial",       callback_data="entrega_obs:parcial")],
        [InlineKeyboardButton("⚠️ Material com avaria",  callback_data="entrega_obs:avaria")],
        [InlineKeyboardButton("🔄 Produto diferente",    callback_data="entrega_obs:diferente")],
        [InlineKeyboardButton("✏️ Outra observação",     callback_data="entrega_obs:outro")],
    ])

def _salvar_entrega_db(pfm_codigo, obs):
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            "UPDATE lancamentos SET obs_entrega=?, "
            "entregue_em=datetime('now','localtime') WHERE pfm_codigo=?",
            (obs, pfm_codigo)
        )

def _tela_apos_entrega(pfm_codigo):
    pedido = buscar_pedido(pfm_codigo)
    if not pedido:
        return None, None
    preparar_visualizacao_pedido(pedido)
    return (
        mostrar_pedido(pedido),
        teclado_pedido(pedido.doc_id, pfm_codigo, pedido.doc_id_nfe,
                       pedido.doc_id_comprovante, pedido.qtd_fotos_entrega, pedido.obs_entrega)
    )

def _adicionar_foto_entrega(pfm_codigo, doc_id_foto, legenda):
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            "INSERT INTO entrega_fotos (pfm_codigo, doc_id, legenda) VALUES (?,?,?)",
            (pfm_codigo, doc_id_foto, legenda)
        )
        qtd = con.execute(
            "SELECT COUNT(*) FROM entrega_fotos WHERE pfm_codigo=?", (pfm_codigo,)
        ).fetchone()[0]
        row = con.execute("SELECT caminho FROM documentos WHERE id=?", (doc_id_foto,)).fetchone()
    caminho_original = row[0] if row else None
    _arquivar_documento(pfm_codigo, f"foto{qtd:02d}", caminho_original, None, _pasta_entrega)

def _listar_fotos_entrega(pfm_codigo):
    with sqlite3.connect(DB_PATH) as con:
        return con.execute(
            "SELECT ef.id, ef.doc_id, ef.legenda, d.caminho FROM entrega_fotos ef "
            "JOIN documentos d ON d.id = ef.doc_id WHERE ef.pfm_codigo=? ORDER BY ef.id",
            (pfm_codigo,)
        ).fetchall()

def _icone_arquivo_entrega(caminho):
    if caminho and Path(caminho).suffix.lower() == ".pdf":
        return "📄"
    return "📷"

def _apagar_foto_entrega(foto_id):
    with sqlite3.connect(DB_PATH) as con:
        con.execute("DELETE FROM entrega_fotos WHERE id=?", (foto_id,))

def _atualizar_obs_entrega(pfm_codigo, obs):
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            "UPDATE lancamentos SET obs_entrega=? WHERE pfm_codigo=?",
            (obs, pfm_codigo)
        )

def _apagar_entrega_db(pfm_codigo):
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            "UPDATE lancamentos SET obs_entrega=NULL, entregue_em=NULL "
            "WHERE pfm_codigo=?", (pfm_codigo,)
        )
        con.execute("DELETE FROM entrega_fotos WHERE pfm_codigo=?", (pfm_codigo,))

def _buscar_estado_entrega(pfm_codigo):
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(
            "SELECT fornecedor, obs_entrega FROM lancamentos WHERE pfm_codigo=?",
            (pfm_codigo,)
        ).fetchone()
        qtd_fotos = con.execute(
            "SELECT COUNT(*) FROM entrega_fotos WHERE pfm_codigo=?", (pfm_codigo,)
        ).fetchone()[0]
    forn, obs = row if row else (None, None)
    return forn, obs, qtd_fotos

def _rotulo_qtd_arquivos(qtd):
    if qtd == 0:
        return "nenhuma"
    if qtd == 1:
        return "1 arquivo"
    return f"{qtd} arquivos"

def _texto_gerir_entrega(pfm_codigo, forn, obs_entrega, qtd_fotos):
    return (
        f"#{pfm_codigo} — {forn or pfm_codigo}\n\n"
        f"Entrega registrada\n"
        f"Observação: {obs_entrega or '—'}\n"
        f"Arquivos: {_rotulo_qtd_arquivos(qtd_fotos)}"
    )

def _teclado_gerir_entrega(pfm_codigo, qtd_fotos):
    botoes = [
        [InlineKeyboardButton("✏️ Mudar observação", callback_data=f"entrega_mudar_obs:{pfm_codigo}")],
    ]
    if qtd_fotos:
        botoes.append([InlineKeyboardButton(f"👀 Ver {_rotulo_qtd_arquivos(qtd_fotos)}", callback_data=f"entrega_ver_fotos:{pfm_codigo}")])
    botoes.append([InlineKeyboardButton("📎 Adicionar foto ou arquivo", callback_data=f"entrega_trocar_foto:{pfm_codigo}")])
    if qtd_fotos:
        botoes.append([InlineKeyboardButton("🗑 Remover arquivo", callback_data=f"entrega_remover_foto:{pfm_codigo}")])
    botoes += [
        [InlineKeyboardButton("❌ Apagar entrega", callback_data=f"entrega_apagar:{pfm_codigo}")],
        [InlineKeyboardButton("← Voltar",          callback_data=f"entrega_voltar:{pfm_codigo}")],
    ]
    return InlineKeyboardMarkup(botoes)

def _teclado_mudar_obs(pfm_codigo):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✅ Entrega completa",      callback_data="entrega_editobs:completa")],
        [InlineKeyboardButton("📦 Entrega parcial",       callback_data="entrega_editobs:parcial")],
        [InlineKeyboardButton("⚠️ Material com avaria",  callback_data="entrega_editobs:avaria")],
        [InlineKeyboardButton("🔄 Produto diferente",    callback_data="entrega_editobs:diferente")],
        [InlineKeyboardButton("✏️ Outra observação",     callback_data="entrega_editobs:outro")],
        [InlineKeyboardButton("← Voltar",                callback_data=f"entrega_editar:{pfm_codigo}")],
    ])

def _texto_obs_entrega(pfm_codigo, forn):
    return f"#{pfm_codigo} — {forn or pfm_codigo}\n\nComo foi a entrega?"

def _mostrar_pedidos_entrega(pedidos):
    if not pedidos:
        return "Nenhum pedido sem entrega registrada."
    linhas = ["Qual pedido chegou?", ""]
    for pfm_codigo, forn, valor, status in pedidos:
        emoji = "🟡" if status == "a_pagar" else "🟢"
        v = f"R$ {_fmt_brl(float(valor))}" if valor else ""
        linhas.append(f"{emoji} #{pfm_codigo} · {forn or '—'}" + (f" · {v}" if v else ""))
    return "\n".join(linhas)

def _teclado_obs_com_cancelar(com_foto=False):
    linhas = list(teclado_obs_entrega().inline_keyboard)
    if not com_foto:
        linhas.append([InlineKeyboardButton("📎 Foto / Documento", callback_data="entrega_foto_primeiro")])
    linhas.append([InlineKeyboardButton("✖ Cancelar", callback_data="entrega_cancelar")])
    return InlineKeyboardMarkup(linhas)

def _tela_categoria(cat, ramo):
    if cat:
        linha_ramo = f"\n{ramo}" if ramo and ramo != "A PREENCHER" else ""
        return f"Como classificar este pedido?\n\n{cat.label()}{linha_ramo}"
    return "Como classificar este pedido?"

def _teclado_selecao_categorias(doc_id, ggv):
    cats = list(CategoriaLancamento)
    botoes = []
    for i in range(0, len(cats), 2):
        linha = [
            InlineKeyboardButton(c.label(), callback_data=f"cat_sel:{doc_id}:{ggv}:{c.value}")
            for c in cats[i:i+2]
        ]
        botoes.append(linha)
    return InlineKeyboardMarkup(botoes)

def _teclado_categoria(doc_id, ggv, cat):
    if cat:
        return InlineKeyboardMarkup([
            [InlineKeyboardButton("✅ Confirmar",   callback_data=f"cat_confirmar:{doc_id}:{ggv}:{cat.value}")],
            [InlineKeyboardButton("Escolher outra", callback_data=f"cat_corrigir:{doc_id}:{ggv}")],
        ])
    return _teclado_selecao_categorias(doc_id, ggv)

def _parse_nfe(corpo: str) -> dict:
    """Extrai campos da NF-e do texto retornado pelo Claude."""
    valor_str = _campo(corpo, "Valor total")
    try:
        valor_v = float(valor_str.replace("R$", "").replace(".", "").replace(",", ".").strip())
    except Exception:
        valor_v = None
    return {
        "numero":    _campo(corpo, "Número da NF"),
        "cnpj":      _campo(corpo, "CNPJ/CPF do emitente"),
        "emitente":  _campo(corpo, "Nome do emitente"),
        "valor_fmt": valor_str,
        "valor_v":   valor_v,
        "data":      _campo(corpo, "Data de emissão"),
        "descricao": _campo(corpo, "Descrição do serviço/produto"),
    }

def _mostrar_nfe(dados: dict, candidatos: list) -> str:
    linhas = ["NF-e identificada.\n"]
    if dados["emitente"] != "A PREENCHER":
        linhas.append(f"{dados['emitente']} — {dados['valor_fmt']}")
    else:
        linhas.append(dados["valor_fmt"])
    if dados["numero"]   != "A PREENCHER": linhas.append(f"NF {dados['numero']}")
    if dados["data"]     != "A PREENCHER": linhas.append(dados["data"])
    if dados["descricao"] != "A PREENCHER": linhas.append(dados["descricao"])
    linhas.append("")
    if not candidatos:
        linhas.append("Nenhum pedido pago sem NF-e encontrado.")
        return "\n".join(linhas)
    fortes = [c for c in candidatos if c["score"] > 0]
    if fortes:
        linhas.append("A qual pedido vincular esta NF-e?\n")
        for c in fortes:
            valor_fmt = f"R$ {_fmt_brl(c['valor_lanc'])}" if c["valor_lanc"] else "—"
            linhas.append(f"🟢 #{c['pfm_codigo']} · {c['fornecedor']} · {valor_fmt}")
    else:
        linhas.append("Escolha o pedido manualmente:\n")
        for c in candidatos:
            valor_fmt = f"R$ {_fmt_brl(c['valor_lanc'])}" if c["valor_lanc"] else "—"
            linhas.append(f"🟢 #{c['pfm_codigo']} · {c['fornecedor']} · {valor_fmt}")
    return "\n".join(linhas)

def _teclado_candidatos_nfe(doc_id: int, candidatos: list):
    botoes = []
    for c in candidatos:
        botoes.append([InlineKeyboardButton(
            f"#{c['pfm_codigo']}",
            callback_data=f"nfe_confirmar:{doc_id}:{c['pfm_codigo']}"
        )])
    botoes.append([InlineKeyboardButton("Nenhum destes", callback_data="nfe_cancelar")])
    return InlineKeyboardMarkup(botoes)

async def _executar_gerar_pfm(query, ctx, doc_id, ggv, categoria):
    await query.edit_message_text("Gerando Pedido de Compra...")
    caminho, codigo, fornecedor, valor_v, lanc_status, ja_existia = gerar_pfm(doc_id, categoria)
    html      = _gerar_html_pc(doc_id)
    pdf_bytes = await _html_para_pdf(html)
    caminho.with_suffix(".pdf").write_bytes(pdf_bytes)
    await ctx.bot.send_document(
        chat_id=DONO_ID,
        document=pdf_bytes,
        filename=f"{codigo}.pdf",
        caption=f"Pedido #{codigo}"
    )
    if ja_existia:
        lanc_msg = f"Pedido #{codigo} já tinha registro financeiro."
    elif lanc_status == "pendente_revisao":
        lanc_msg = (
            f"🔴 Pedido #{codigo} requer atenção.\n"
            f"Fornecedor ou valor ausente — verifique antes de pagar."
        )
    else:
        cat_linha = f"\n{categoria.label()}" if categoria else ""
        lanc_msg = (
            f"🟡 {codigo} — aguardando pagamento\n\n"
            f"{fornecedor} — R$ {_fmt_brl(valor_v)}{cat_linha}"
        )
    await ctx.bot.send_message(chat_id=DONO_ID, text=lanc_msg)


async def _executar_revisao_pfm(query, ctx, doc_id, pfm_codigo_base):
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute("SELECT rev_numero FROM documentos WHERE id=?", (doc_id,)).fetchone()
        rev_num = (row[0] or 0) + 1
        con.execute("UPDATE documentos SET rev_numero=? WHERE id=?", (rev_num, doc_id))
    rev_codigo = f"{pfm_codigo_base}-R{rev_num:02d}"
    await query.edit_message_text(f"Gerando revisão {rev_codigo}...")
    caminho_rev, *_ = gerar_pfm(doc_id, pfm_codigo_override=rev_codigo)
    # Sobrescreve o DOCX principal para manter OneDrive sempre atualizado
    nome_principal    = caminho_rev.name.replace(rev_codigo, pfm_codigo_base, 1)
    caminho_principal = caminho_rev.parent / nome_principal
    shutil.copy2(caminho_rev, caminho_principal)
    atualizar(doc_id, caminho_pfm=str(caminho_principal))
    html      = _gerar_html_pc(doc_id)
    pdf_bytes = await _html_para_pdf(html)
    caminho_rev.with_suffix(".pdf").write_bytes(pdf_bytes)
    caminho_principal.with_suffix(".pdf").write_bytes(pdf_bytes)
    await ctx.bot.send_document(
        chat_id=DONO_ID,
        document=pdf_bytes,
        filename=f"{rev_codigo}.pdf",
        caption=f"📄 {rev_codigo}"
    )
    ctx.user_data.pop("modo_revisao", None)
    await ctx.bot.send_message(
        chat_id=DONO_ID,
        text=f"✅ {rev_codigo} gerado. Lançamento financeiro mantido.",
        reply_markup=teclado_pedido(doc_id, pfm_codigo_base)
    )

async def _sincronizar_receita_pendentes(ctx: ContextTypes.DEFAULT_TYPE):
    """Job periódico: tenta de novo os fornecedores que ficaram sem resposta da Receita na hora do cadastro."""
    with sqlite3.connect(DB_PATH) as con:
        pendentes = con.execute(
            "SELECT id, cnpj FROM fornecedores WHERE receita_pendente=1"
        ).fetchall()
    if not pendentes:
        return

    resolvidos = 0
    for forn_id, cnpj in pendentes:
        cnpj_digits = re.sub(r"\D", "", cnpj or "")
        if len(cnpj_digits) != 14:
            continue
        receita = _consultar_receita(cnpj_digits)
        if receita:
            with sqlite3.connect(DB_PATH) as con:
                con.execute(
                    "UPDATE fornecedores SET razao_social=COALESCE(razao_social,?), "
                    "cidade=COALESCE(cidade,?), uf=COALESCE(uf,?), receita_pendente=0 WHERE id=?",
                    (receita["razao_social"], receita["cidade"], receita["uf"], forn_id)
                )
            resolvidos += 1

    total = len(pendentes)
    if resolvidos == total:
        texto = f"📋 Receita sincronizada — {resolvidos} de {total} pendências resolvidas."
    else:
        restantes = total - resolvidos
        verbo = "segue" if restantes == 1 else "seguem"
        texto = (f"📋 Receita sincronizada — {resolvidos} de {total} pendências resolvidas. "
                 f"{restantes} {verbo} tentando.")
    await ctx.bot.send_message(chat_id=DONO_ID, text=texto)


# ── Handlers ───────────────────────────────────────────────────────────────

async def start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    if TEST_MODE:
        await update.message.reply_text(
            "🧪 MODO TESTE ATIVO\n"
            "Banco: data/laura_test.db\n"
            "Uploads: data/test_uploads\n"
            "Pedidos: data/test_pfms"
        )
    else:
        await update.message.reply_text("Estou online.")

async def ajuda(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    await update.message.reply_text(
        mostrar_ajuda(),
        reply_markup=teclado_ajuda(),
        parse_mode="HTML"
    )

async def comando_desconhecido(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    await update.message.reply_text(
        mostrar_boas_vindas(),
        reply_markup=teclado_boas_vindas()
    )

async def obras_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    obras = _listar_obras()
    await update.message.reply_text(
        mostrar_lista_obras(obras),
        reply_markup=teclado_lista_obras(obras)
    )

async def entrega_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    pedidos = buscar_pedidos_sem_entrega()
    if not pedidos:
        await update.message.reply_text("Nenhum pedido sem entrega registrada.")
        return
    ctx.user_data["entrega_doc_id_foto"] = None
    if len(pedidos) == 1:
        pfm_codigo, forn, _, _ = pedidos[0]
        ctx.user_data["entrega_pfm_codigo"] = pfm_codigo
        await update.message.reply_text(
            _texto_obs_entrega(pfm_codigo, forn),
            reply_markup=_teclado_obs_com_cancelar()
        )
    else:
        await update.message.reply_text(
            _mostrar_pedidos_entrega(pedidos),
            reply_markup=_teclado_pedidos_entrega(pedidos)
        )

async def nova_obra(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    ctx.user_data["aguardando"] = "nova_obra_codigo"
    await update.message.reply_text("Código da nova obra (ex: GGV04):")

async def receber_arquivo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return

    if update.message.photo:
        tg_file = await update.message.photo[-1].get_file()
        nome = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
        mime = None
    elif update.message.document:
        tg_file = await update.message.document.get_file()
        nome = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{update.message.document.file_name}"
        mime = update.message.document.mime_type or "application/pdf"
    else:
        await update.message.reply_text("Recebi. Ainda não sei o que fazer com isso.")
        return

    conteudo = await tg_file.download_as_bytearray()
    hash_arquivo = hashlib.sha256(conteudo).hexdigest()
    if TEST_MODE:
        hash_arquivo += f"_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"

    if ja_existe(hash_arquivo):
        await update.message.reply_text("Este arquivo já foi recebido.")
        return

    if mime is None:
        mime = "image/png" if conteudo[:4] == b'\x89PNG' else "image/jpeg"

    ACEITOS = {"image/jpeg", "image/png", "image/gif", "image/webp", "application/pdf"}
    if mime not in ACEITOS:
        await update.message.reply_text(
            f"Formato não suportado: {mime}\n\n"
            "Aceito: foto (JPEG, PNG, GIF, WEBP) ou PDF."
        )
        return

    caminho = UPLOADS / nome
    caminho.write_bytes(conteudo)
    if TEST_MODE:
        await update.message.reply_text("🧪 MODO TESTE ATIVO — dados não afetam produção.")

    doc_id = registrar(nome, caminho, hash_arquivo, "pendente", "nao_identificado", "")

    if ctx.user_data.get("aguardando") == "foto_entrega_obs":
        ctx.user_data["aguardando"] = "entrega_legenda_inicial"
        atualizar(doc_id, tipo="foto_entrega")
        ctx.user_data["entrega_doc_id_foto"] = doc_id
        await update.message.reply_text("Legenda do arquivo (ex: nota fiscal, caixa avariada):")
        return

    if ctx.user_data.get("aguardando") == "foto_entrega_troca":
        ctx.user_data["aguardando"] = "entrega_legenda_add"
        atualizar(doc_id, tipo="foto_entrega")
        ctx.user_data["entrega_doc_id_foto"] = doc_id
        await update.message.reply_text("Legenda do arquivo (ex: nota fiscal, caixa avariada):")
        return

    await update.message.reply_text(
        "O que é este documento?",
        reply_markup=teclado_tipo_inicial(doc_id)
    )

async def receber_texto(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != DONO_ID:
        return
    aguardando = ctx.user_data.get("aguardando")
    texto      = update.message.text.strip()

    if not aguardando:
        if _SAUDACOES_RE.match(texto):
            await update.message.reply_text(
                mostrar_boas_vindas(),
                reply_markup=teclado_boas_vindas()
            )
            return
        if _OBRAS_RE.match(texto):
            obras = _listar_obras()
            await update.message.reply_text(
                mostrar_lista_obras(obras),
                reply_markup=teclado_lista_obras(obras)
            )
            return
        m_obra = GGV_CODIGO_RE.match(texto)
        if m_obra:
            codigo = m_obra.group(1).upper()
            obra = buscar_obra(codigo)
            if obra:
                await update.message.reply_text(
                    mostrar_cockpit_obra(obra),
                    reply_markup=teclado_obra(codigo, _pedidos_obra(codigo))
                )
            else:
                await update.message.reply_text(f"Obra {codigo} não encontrada.")
            return
        m = PFM_CODIGO_RE.search(texto)
        if m:
            pfm_codigo = m.group(1).upper()
            pedido = buscar_pedido(pfm_codigo)
            if pedido:
                preparar_visualizacao_pedido(pedido)
                await update.message.reply_text(
                    mostrar_pedido(pedido),
                    reply_markup=teclado_pedido(pedido.doc_id, pfm_codigo, pedido.doc_id_nfe, pedido.doc_id_comprovante, pedido.qtd_fotos_entrega, pedido.obs_entrega)
                )
            else:
                await update.message.reply_text(f"Pedido {pfm_codigo} não encontrado.")
        else:
            await update.message.reply_text(
                mostrar_boas_vindas(),
                reply_markup=teclado_boas_vindas()
            )
        return

    doc_id = ctx.user_data.get("doc_id")
    ggv    = ctx.user_data.get("ggv")

    if aguardando == "condicao_pgto":
        atualizar(doc_id, condicao_pgto=texto)
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando == "data_entrega":
        atualizar(doc_id, data_entrega=texto)
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando == "edit_desconto":
        try:
            v = _parse_brl(re.sub(r"[^\d,.]", "", texto))
        except Exception:
            v = 0.0
        atualizar(doc_id, desconto_rs=f"{v:.2f}")
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando == "obs_entrega_texto":
        obs = texto.strip()
        pfm_codigo = ctx.user_data.pop("entrega_pfm_codigo", None)
        doc_id_foto = ctx.user_data.pop("entrega_doc_id_foto", None)
        legenda_foto = ctx.user_data.pop("entrega_legenda_foto", None)
        ctx.user_data["aguardando"] = None
        if pfm_codigo:
            _salvar_entrega_db(pfm_codigo, obs)
            if doc_id_foto:
                _adicionar_foto_entrega(pfm_codigo, doc_id_foto, legenda_foto)
            texto_ped, markup = _tela_apos_entrega(pfm_codigo)
            await update.message.reply_text(texto_ped or "Entrega registrada.", reply_markup=markup)
        else:
            await update.message.reply_text("Pedido não encontrado.")

    elif aguardando == "edit_obs_entrega_texto":
        obs = texto.strip()
        pfm_codigo = ctx.user_data.pop("entrega_pfm_codigo", None)
        ctx.user_data["aguardando"] = None
        if pfm_codigo:
            _atualizar_obs_entrega(pfm_codigo, obs)
            forn, obs_ent, qtd_fotos = _buscar_estado_entrega(pfm_codigo)
            await update.message.reply_text(
                _texto_gerir_entrega(pfm_codigo, forn, obs_ent, qtd_fotos),
                reply_markup=_teclado_gerir_entrega(pfm_codigo, qtd_fotos)
            )
        else:
            await update.message.reply_text("Pedido não encontrado.")

    elif aguardando == "entrega_legenda_inicial":
        legenda = texto.strip()
        ctx.user_data["entrega_legenda_foto"] = legenda
        ctx.user_data["aguardando"] = None
        pfm_codigo = ctx.user_data.get("entrega_pfm_codigo")
        forn, _, _ = _buscar_estado_entrega(pfm_codigo) if pfm_codigo else (None, None, None)
        await update.message.reply_text(
            _texto_obs_entrega(pfm_codigo, forn),
            reply_markup=_teclado_obs_com_cancelar(com_foto=True)
        )

    elif aguardando == "entrega_legenda_add":
        legenda = texto.strip()
        pfm_codigo = ctx.user_data.pop("entrega_pfm_codigo", None)
        doc_id_foto = ctx.user_data.pop("entrega_doc_id_foto", None)
        ctx.user_data["aguardando"] = None
        if pfm_codigo and doc_id_foto:
            _adicionar_foto_entrega(pfm_codigo, doc_id_foto, legenda)
            forn, obs_ent, qtd_fotos = _buscar_estado_entrega(pfm_codigo)
            await update.message.reply_text(
                _texto_gerir_entrega(pfm_codigo, forn, obs_ent, qtd_fotos),
                reply_markup=_teclado_gerir_entrega(pfm_codigo, qtd_fotos)
            )
        else:
            await update.message.reply_text("Pedido não encontrado.")

    elif aguardando == "edit_contato":
        m_fone = re.search(r'\s+([\d][\d\s\(\)\-\.]{5,})\s*$', texto)
        if m_fone:
            nome_v = texto[:m_fone.start()].strip()
            fone_v = m_fone.group(1).strip()
        else:
            nome_v = texto.strip()
            fone_v = ""
        with sqlite3.connect(DB_PATH) as con:
            row = con.execute("SELECT dados_claude FROM documentos WHERE id=?", (doc_id,)).fetchone()
        if row:
            novos_dados = _substituir_campo(row[0], "Vendedor", nome_v)
            if fone_v:
                novos_dados = _substituir_campo(novos_dados, "Telefone do vendedor", fone_v)
            atualizar(doc_id, dados_claude=novos_dados)
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando in ("edit_fornecedor", "edit_cnpj", "edit_valor", "edit_pix", "edit_itens"):
        campo_map = {
            "edit_fornecedor": "Fornecedor",
            "edit_cnpj":       "CNPJ/CPF",
            "edit_valor":      "Valor total",
            "edit_pix":        "Chave PIX",
        }
        tipo = ctx.user_data.get("tipo", "orcamento")
        with sqlite3.connect(DB_PATH) as con:
            row = con.execute(
                "SELECT dados_claude, ggv FROM documentos WHERE id=?", (doc_id,)
            ).fetchone()
        if row:
            dados_atuais, ggv_db = row
            if aguardando == "edit_itens":
                novos_dados = _recalcular_itens(_substituir_itens(dados_atuais, texto))
                nome_campo = "Itens"
            else:
                nome_campo = campo_map[aguardando]
                novos_dados = _substituir_campo(dados_atuais, nome_campo, texto)
            atualizar(doc_id, dados_claude=novos_dados)
            ctx.user_data["aguardando"] = None
            texto_resumo, markup = _resumo_gerar(doc_id)
            await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")
        else:
            ctx.user_data["aguardando"] = None
            await update.message.reply_text("Documento não encontrado.")

    elif aguardando == "endereco_entrega":
        atualizar(doc_id, endereco_entrega=texto)
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando == "vencimento_pgto":
        atualizar(doc_id, vencimento_pgto=texto)
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando == "edit_encarregado":
        atualizar(doc_id, encarregado=texto)
        ctx.user_data["aguardando"] = None
        texto_resumo, markup = _resumo_gerar(doc_id)
        await update.message.reply_text(texto_resumo, reply_markup=markup, parse_mode="HTML")

    elif aguardando == "nova_obra_codigo":
        codigo = texto.upper()
        if not GGV_CODIGO_RE.match(codigo):
            await update.message.reply_text("Formato inválido. Use GGV seguido de dois dígitos — ex: GGV04. Tente novamente:")
            return
        criada = criar_obra(codigo)
        ctx.user_data["aguardando"] = None
        if not criada:
            await update.message.reply_text(f"Obra {codigo} já existe.")
        else:
            await update.message.reply_text(f"Obra {codigo} criada. Complete os dados abaixo.")
        obra = buscar_obra(codigo)
        await update.message.reply_text(mostrar_cockpit_obra(obra), reply_markup=teclado_obra(codigo, _pedidos_obra(codigo)))

    elif aguardando and aguardando.startswith("obra_edit_"):
        campo = aguardando[len("obra_edit_"):]
        obra_codigo = ctx.user_data.get("obra_codigo")
        if obra_codigo and campo:
            atualizar_obra(obra_codigo, **{campo: texto})
            ctx.user_data["aguardando"] = None
            obra = buscar_obra(obra_codigo)
            await update.message.reply_text(
                mostrar_cockpit_obra(obra),
                reply_markup=teclado_obra(obra_codigo, _pedidos_obra(obra_codigo))
            )
        else:
            ctx.user_data["aguardando"] = None
            await update.message.reply_text("Contexto perdido. Envie o código da obra novamente.")

async def responder_botao(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    partes = query.data.split(":")
    acao   = partes[0]

    # GGV blocking — alert antes de responder, depois retorna
    if acao == "ok" and partes[3] == "nao_identificado":
        try:
            await query.answer("Selecione a obra antes de confirmar.", show_alert=True)
        except Exception:
            pass
        return

    try:
        await query.answer()
    except Exception:
        return  # callback expirado (bot reiniciado) — ignora silenciosamente

    try:
        if acao == "ok":
            _, doc_id, tipo, ggv = partes
            atualizar(int(doc_id), status="confirmado")
            if tipo == "orcamento":
                ctx.user_data.update({"doc_id": int(doc_id), "ggv": ggv, "aguardando": None})
                texto, markup = _resumo_gerar(int(doc_id))
                await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")
            elif tipo == "comprovante_pix":
                dados_claude = _dados_doc(int(doc_id))
                dados        = parse_comprovante(dados_claude)
                candidatos   = buscar_candidatos_pix(dados["valor_v"], dados["favorecido"], dados["cnpj"])
                await query.edit_message_text(mostrar_comprovante_candidatos(dados, candidatos))
            else:
                emoji, label = TIPOS.get(tipo, ("📄", tipo))
                await query.edit_message_text(f"Confirmado: {label}")

        elif acao == "cancelar":
            atualizar(int(partes[1]), status="cancelado")
            await query.edit_message_text("Cancelado.")

        elif acao == "sel_tipo":
            _, doc_id, tipo, ggv = partes
            botoes = [[InlineKeyboardButton(f"{e} {l}", callback_data=f"set_tipo:{doc_id}:{t}:{ggv}")]
                      for t, (e, l) in TIPOS.items()]
            await query.edit_message_reply_markup(InlineKeyboardMarkup(botoes))

        elif acao == "set_tipo":
            _, doc_id, novo_tipo, ggv = partes
            atualizar(int(doc_id), tipo=novo_tipo)
            with sqlite3.connect(DB_PATH) as con:
                status = con.execute("SELECT status FROM documentos WHERE id=?", (int(doc_id),)).fetchone()[0]
            if status == "confirmado":
                texto, markup = _resumo_gerar(int(doc_id))
                await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")
            else:
                texto, markup = _resumo_gerar(int(doc_id))
                await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")

        elif acao == "sel_tipo_inicial":
            _, doc_id, tipo = partes
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute("SELECT caminho FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
            if not row:
                await query.edit_message_text("Documento não encontrado.")
                return
            caminho_doc = row[0]

            if tipo == "foto_entrega":
                atualizar(int(doc_id), tipo=tipo)
                pedidos = buscar_pedidos_sem_entrega()
                if not pedidos:
                    await query.edit_message_text("Nenhum pedido sem entrega registrada.")
                    return
                ctx.user_data["entrega_doc_id_foto"] = int(doc_id)
                await query.edit_message_text(
                    _mostrar_pedidos_entrega(pedidos),
                    reply_markup=_teclado_pedidos_entrega(pedidos)
                )
                return

            mime_inf = "application/pdf" if caminho_doc.lower().endswith(".pdf") else "image/jpeg"
            tipo_conteudo = "document" if mime_inf == "application/pdf" else "image"
            conteudo = Path(caminho_doc).read_bytes()
            dados_b64 = base64.standard_b64encode(conteudo).decode()
            await query.edit_message_text("Analisando...")
            resposta = claude.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=4096,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": tipo_conteudo, "source": {"type": "base64", "media_type": mime_inf, "data": dados_b64}},
                        {"type": "text", "text": PROMPT}
                    ]
                }]
            )
            _, ggv, corpo = parse_resposta(resposta.content[0].text)
            atualizar(int(doc_id), tipo=tipo, ggv=ggv, dados_claude=corpo)
            emoji, label_tipo = TIPOS.get(tipo, ("📄", tipo))
            label_ggv = ggv if ggv != "nao_identificado" else "Obra não identificada"
            if tipo == "comprovante_pix":
                dados = parse_comprovante(corpo)
                ident = dados["id_transacao"] if dados["id_transacao"] != "A PREENCHER" else None
                ja_pago = None
                if ident and not TEST_MODE:
                    with sqlite3.connect(DB_PATH) as con:
                        ja_pago = con.execute(
                            "SELECT pfm_codigo FROM lancamentos "
                            "WHERE identificador_comprovante=? AND status='pago' LIMIT 1",
                            (ident,)
                        ).fetchone()
                if ja_pago:
                    await query.edit_message_text(
                        f"Comprovante já registrado.\n\n"
                        f"Pedido #{ja_pago[0]} — 🟢 Pago\n\n"
                        "Cada comprovante pode ser usado apenas uma vez."
                    )
                else:
                    candidatos = buscar_candidatos_pix(dados["valor_v"], dados["favorecido"], dados["cnpj"])
                    markup     = teclado_candidatos_pix(int(doc_id), candidatos) if candidatos else None
                    await query.edit_message_text(
                        mostrar_comprovante_candidatos(dados, candidatos),
                        reply_markup=markup
                    )
            elif tipo == "nota_fiscal":
                dados_nfe = _parse_nfe(corpo)
                candidatos = buscar_candidatos_nfe(dados_nfe["cnpj"], dados_nfe["valor_v"], DB_PATH)
                await query.edit_message_text(
                    _mostrar_nfe(dados_nfe, candidatos),
                    reply_markup=_teclado_candidatos_nfe(int(doc_id), candidatos)
                )
            elif tipo == "orcamento":
                texto, markup = _resumo_gerar(int(doc_id))
                await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")
            else:
                await query.edit_message_text(
                    f"{label_tipo}\n\n{corpo}\n\nConfirmar ou cancelar?",
                    reply_markup=InlineKeyboardMarkup([[
                        InlineKeyboardButton("✅ Confirmar", callback_data=f"ok:{doc_id}:{tipo}:{ggv}"),
                        InlineKeyboardButton("Cancelar",    callback_data=f"cancelar:{doc_id}"),
                    ]])
                )

        elif acao == "pix_confirmar":
            _, doc_id_comp, pfm_codigo = partes
            with sqlite3.connect(DB_PATH) as con:
                comp = con.execute(
                    "SELECT dados_claude FROM documentos WHERE id=?", (int(doc_id_comp),)
                ).fetchone()
                lanc = con.execute(
                    "SELECT fornecedor, valor, status FROM lancamentos WHERE pfm_codigo=?",
                    (pfm_codigo,)
                ).fetchone()
            if not comp or not lanc:
                await query.edit_message_text("Dados não encontrados.")
                return
            dados_comp   = parse_comprovante(comp[0])
            forn, valor_lanc, status_lanc = lanc
            status_label = {"a_pagar": "🟡 Aguardando pagamento", "pago": "🟢 Pago"}.get(status_lanc, status_lanc)
            valor_lanc_fmt = f"R$ {_fmt_brl(valor_lanc)}" if valor_lanc else "—"
            texto = (
                f"Confirmar pagamento?\n\n"
                f"Comprovante:  R$ {dados_comp['valor_fmt']}  {dados_comp['data']}\n"
                f"Pedido:       #{pfm_codigo} — {forn}\n"
                f"Valor:        {valor_lanc_fmt}\n"
                f"Status:       {status_label}"
            )
            await query.edit_message_text(texto, reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("✅ Confirmar pagamento",
                                     callback_data=f"pix_pagar:{doc_id_comp}:{pfm_codigo}")],
                [InlineKeyboardButton("↩️ Voltar",
                                     callback_data="pix_cancelar")],
            ]))

        elif acao == "pix_pagar":
            _, doc_id_comp, pfm_codigo = partes
            with sqlite3.connect(DB_PATH) as con:
                comp = con.execute(
                    "SELECT dados_claude FROM documentos WHERE id=?", (int(doc_id_comp),)
                ).fetchone()
            if not comp:
                await query.edit_message_text("Comprovante não encontrado.")
                return
            dados_comp = parse_comprovante(comp[0])
            ident_comp = dados_comp["id_transacao"] if dados_comp["id_transacao"] != "A PREENCHER" else None
            ja_usado = None
            if ident_comp and not TEST_MODE:
                with sqlite3.connect(DB_PATH) as con:
                    ja_usado = con.execute(
                        "SELECT pfm_codigo FROM lancamentos "
                        "WHERE identificador_comprovante=? AND status='pago' LIMIT 1",
                        (ident_comp,)
                    ).fetchone()
            if ja_usado:
                await query.edit_message_text(
                    f"Comprovante já registrado no Pedido #{ja_usado[0]}.\n"
                    "Cada comprovante pode ser usado apenas uma vez."
                )
                return
            _MESES = {"janeiro":"01","fevereiro":"02","março":"03","abril":"04",
                      "maio":"05","junho":"06","julho":"07","agosto":"08",
                      "setembro":"09","outubro":"10","novembro":"11","dezembro":"12"}
            _data_raw = dados_comp["data"] if dados_comp["data"] != "A PREENCHER" \
                        else datetime.now().strftime("%d/%m/%Y")
            data_pgto = _data_raw
            for nome, num in _MESES.items():
                if nome in _data_raw.lower():
                    data_pgto = re.sub(nome, num, _data_raw, flags=re.IGNORECASE)
                    break
            with sqlite3.connect(DB_PATH) as con:
                cur = con.execute(
                    """UPDATE lancamentos
                       SET status='pago', valor_pago=?, data_pagamento=?,
                           doc_id_comprovante=?, identificador_comprovante=?
                       WHERE pfm_codigo=? AND status='a_pagar'""",
                    (dados_comp["valor_v"] or None, data_pgto,
                     int(doc_id_comp), ident_comp, pfm_codigo)
                )
                rowcount = cur.rowcount
            if rowcount == 0:
                await query.edit_message_text(
                    "Não foi possível registrar o pagamento.\n"
                    "O pedido pode já estar pago ou ter sido alterado."
                )
                return
            with sqlite3.connect(DB_PATH) as con:
                row_comp = con.execute("SELECT caminho FROM documentos WHERE id=?", (int(doc_id_comp),)).fetchone()
                row_lanc = con.execute(
                    "SELECT categoria, doc_id FROM lancamentos WHERE pfm_codigo=?", (pfm_codigo,)
                ).fetchone()
            _arquivar_doc_financeiro(pfm_codigo, "comprovante", row_comp[0] if row_comp else None, data_pgto)
            if row_lanc and row_lanc[0] in CATEGORIAS_SEM_NFE_OBRIGATORIA:
                # Taxa/imposto/serviço público: a fatura já enviada é a "terceira via" — não há NF-e separada
                with sqlite3.connect(DB_PATH) as con:
                    row_fatura = con.execute("SELECT caminho FROM documentos WHERE id=?", (row_lanc[1],)).fetchone()
                _arquivar_doc_financeiro(pfm_codigo, "fatura", row_fatura[0] if row_fatura else None, data_pgto)
            await query.edit_message_text(
                f"🟢 Pedido #{pfm_codigo} — pago.\n\n"
                "Envie a NF-e para fechar este pedido."
            )

        elif acao == "pix_cancelar":
            await query.edit_message_text("Cancelado.")

        elif acao == "sel_ggv":
            _, doc_id, tipo, ggv = partes
            botoes = [[InlineKeyboardButton(g, callback_data=f"set_ggv:{doc_id}:{tipo}:{g}")] for g in GGVS]
            botoes.append([InlineKeyboardButton("❓ Não identificado",
                                                callback_data=f"set_ggv:{doc_id}:{tipo}:nao_identificado")])
            botoes.append([InlineKeyboardButton("◀️ Voltar",
                                                callback_data=f"voltar_edit:{doc_id}:{tipo}:{ggv}")])
            await query.edit_message_reply_markup(InlineKeyboardMarkup(botoes))

        elif acao == "set_ggv":
            _, doc_id, tipo, novo_ggv = partes
            atualizar(int(doc_id), ggv=novo_ggv)
            with sqlite3.connect(DB_PATH) as con:
                status = con.execute("SELECT status FROM documentos WHERE id=?", (int(doc_id),)).fetchone()[0]
            if status == "confirmado":
                texto, markup = _resumo_gerar(int(doc_id))
                await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")
            else:
                texto, markup = _resumo_gerar(int(doc_id))
                await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")

        elif acao == "pgto":
            _, doc_id, ggv, escolha = partes
            if escolha == "outro":
                ctx.user_data.update({"doc_id": int(doc_id), "ggv": ggv, "aguardando": "condicao_pgto"})
                await query.edit_message_text("Condição de pagamento:")
                return
            label_pgto = CONDICOES.get(escolha, escolha)
            atualizar(int(doc_id), condicao_pgto=label_pgto)
            ctx.user_data.update({"doc_id": int(doc_id), "ggv": ggv, "aguardando": None})
            texto, markup = _resumo_gerar(int(doc_id))
            await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")

        elif acao == "end":
            _, doc_id, ggv, escolha = partes
            if escolha == "outro":
                ctx.user_data.update({"doc_id": int(doc_id), "ggv": ggv, "aguardando": "endereco_entrega"})
                await query.edit_message_text("Endereço de entrega:")
                return
            if escolha.startswith("obra_"):
                ggv_key = escolha[5:]
                endereco = buscar_obra(ggv_key).get("endereco_entrega") or ENDERECOS.get(escolha, escolha)
            else:
                endereco = ENDERECOS.get(escolha, escolha)
            atualizar(int(doc_id), endereco_entrega=endereco)
            texto, markup = _resumo_gerar(int(doc_id))
            await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")

        elif acao == "sel_edit":
            _, doc_id, tipo, ggv = partes
            botoes = [
                [InlineKeyboardButton("👤 Fornecedor",    callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:fornecedor")],
                [InlineKeyboardButton("🔢 CNPJ/CPF",      callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:cnpj")],
                [InlineKeyboardButton("💲 Valor total",   callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:valor")],
                [InlineKeyboardButton("🔑 Chave PIX",     callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:pix")],
                [InlineKeyboardButton("📦 Itens",         callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:itens")],
                [InlineKeyboardButton("🏷️ Desconto",     callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:desconto")],
                [InlineKeyboardButton("💰 Condição pgto", callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:pgto")],
                [InlineKeyboardButton("📅 Data entrega",  callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:data")],
                [InlineKeyboardButton("📍 Endereço",      callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:endereco")],
                [InlineKeyboardButton("📅 Vencimento pgto", callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:vencimento")],
                [InlineKeyboardButton("👷 Encarregado",     callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:encarregado")],
                [InlineKeyboardButton("📞 Contato vendedor", callback_data=f"edit_campo:{doc_id}:{tipo}:{ggv}:contato")],
                [InlineKeyboardButton("🏗 GGV",             callback_data=f"sel_ggv:{doc_id}:{tipo}:{ggv}")],
                [InlineKeyboardButton("📋 Tipo doc.",      callback_data=f"sel_tipo:{doc_id}:{tipo}:{ggv}")],
                [InlineKeyboardButton("◀️ Voltar",         callback_data=f"voltar_edit:{doc_id}:{tipo}:{ggv}")],
            ]
            await query.edit_message_reply_markup(InlineKeyboardMarkup(botoes))

        elif acao == "edit_campo":
            _, doc_id, tipo, ggv, campo = partes
            ctx.user_data.update({"doc_id": int(doc_id), "ggv": ggv, "tipo": tipo})

            if campo == "pgto":
                ctx.user_data["aguardando"] = None
                await query.edit_message_text(
                    "Condição de pagamento:",
                    reply_markup=teclado_condicao(doc_id, tipo, ggv)
                )
            elif campo == "endereco":
                ctx.user_data["aguardando"] = None
                await query.edit_message_text(
                    "Endereço de entrega:",
                    reply_markup=teclado_endereco(doc_id, tipo, ggv)
                )
            elif campo == "vencimento":
                ctx.user_data["aguardando"] = "vencimento_pgto"
                with sqlite3.connect(DB_PATH) as con:
                    row = con.execute("SELECT vencimento_pgto FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
                atual = row[0] if row and row[0] else "Não informado"
                await query.edit_message_text(
                    f"Atual: {atual}\n\nNovo vencimento (ex: Parcela única até 15/07/2026):"
                )
            elif campo == "encarregado":
                ctx.user_data["aguardando"] = "edit_encarregado"
                with sqlite3.connect(DB_PATH) as con:
                    row = con.execute("SELECT encarregado FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
                atual = (row[0] if row and row[0] else None) or buscar_obra(ggv).get("encarregado_nome", "Não definido")
                await query.edit_message_text(
                    f"Atual: {atual}\n\nNovo encarregado:"
                )
            elif campo == "contato":
                ctx.user_data["aguardando"] = "edit_contato"
                with sqlite3.connect(DB_PATH) as con:
                    row = con.execute("SELECT dados_claude FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
                dados_atuais = row[0] if row else ""
                nome_v = _campo(dados_atuais, "Vendedor")
                fone_v = _campo(dados_atuais, "Telefone do vendedor")
                if nome_v == "A PREENCHER": nome_v = "Não informado"
                if fone_v == "A PREENCHER": fone_v = ""
                atual = f"{nome_v}  {fone_v}".strip() if fone_v else nome_v
                await query.edit_message_text(
                    f"Atual: {atual}\n\nNome e telefone do vendedor (ex: Flávio 42 99912-7781):"
                )
            elif campo == "itens":
                ctx.user_data["aguardando"] = "edit_itens"
                with sqlite3.connect(DB_PATH) as con:
                    row = con.execute("SELECT dados_claude FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
                bloco = _bloco_itens(row[0] if row else "")
                await query.edit_message_text(
                    f"Itens atuais:\n\n{bloco}\n\n"
                    "Novos itens:\n"
                    "Use o formato para cálculo automático:\n"
                    "1. Descrição (qtde UN) — R$ valor_total\n"
                    "Ex: 1. Cimento 50kg (10 sc) — R$ 350,00"
                )
            else:
                aguardando_campo = "data_entrega" if campo == "data" else f"edit_{campo}"
                ctx.user_data["aguardando"] = aguardando_campo
                labels = {
                    "fornecedor": "nome do fornecedor",
                    "cnpj":       "CNPJ ou CPF",
                    "valor":      "valor total (ex: R$ 1.500,00)",
                    "pix":        "chave PIX",
                    "desconto":   "valor do desconto em R$ (ex: 80 ou 80,00) — ou 0 para remover",
                    "data":       "data de entrega (ex: 01/08/2026, 7 dias úteis, A combinar)",
                }
                campo_doc = {
                    "fornecedor": "Fornecedor",
                    "cnpj":       "CNPJ/CPF",
                    "valor":      "Valor total",
                    "pix":        "Chave PIX",
                }
                with sqlite3.connect(DB_PATH) as con:
                    row = con.execute(
                        "SELECT dados_claude, desconto_rs, data_entrega FROM documentos WHERE id=?", (int(doc_id),)
                    ).fetchone()
                dados_atuais, desconto_atual, data_atual = row if row else ("", None, None)
                if campo == "desconto":
                    atual = f"R$ {_fmt_brl(_parse_brl(re.sub(r'[^\d,.]', '', str(desconto_atual))))}" if desconto_atual else "Não informado"
                elif campo == "data":
                    atual = data_atual or "Não informada"
                else:
                    atual = _campo(dados_atuais, campo_doc.get(campo, campo))
                await query.edit_message_text(
                    f"Atual: {atual}\n\nNovo valor:"
                )

        elif acao == "voltar_edit":
            _, doc_id, tipo, ggv = partes
            texto, markup = _resumo_gerar(int(doc_id))
            await query.edit_message_text(texto, reply_markup=markup, parse_mode="HTML")

        elif acao == "ver_itens":
            _, doc_id, tipo, ggv = partes
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute("SELECT dados_claude FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
            bloco = _bloco_itens(row[0] if row else "")
            await query.edit_message_text(
                bloco,
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("◀️ Voltar ao resumo", callback_data=f"voltar_edit:{doc_id}:{tipo}:{ggv}")]
                ])
            )

        elif acao == "pfm":
            _, doc_id, ggv = partes
            pfm_codigo_revisao = ctx.user_data.get("modo_revisao")
            if pfm_codigo_revisao:
                await _executar_revisao_pfm(query, ctx, int(doc_id), pfm_codigo_revisao)
            else:
                atualizar(int(doc_id), status="confirmado")
                ramo = _campo(_dados_doc(int(doc_id)), "Ramo de atividade")
                cat  = sugerir_categoria(ramo)
                await query.edit_message_text(
                    _tela_categoria(cat, ramo),
                    reply_markup=_teclado_categoria(int(doc_id), ggv, cat)
                )

        elif acao == "cat_confirmar":
            _, doc_id, ggv, cat_val = partes
            await _executar_gerar_pfm(query, ctx, int(doc_id), ggv, CategoriaLancamento(cat_val))

        elif acao == "cat_corrigir":
            _, doc_id, ggv = partes
            await query.edit_message_text(
                "Selecione a categoria:",
                reply_markup=_teclado_selecao_categorias(int(doc_id), ggv)
            )

        elif acao == "cat_sel":
            _, doc_id, ggv, cat_val = partes
            await _executar_gerar_pfm(query, ctx, int(doc_id), ggv, CategoriaLancamento(cat_val))

        elif acao == "pfm_revisar":
            _, doc_id, pfm_codigo = partes
            ctx.user_data["doc_id"]       = int(doc_id)
            ctx.user_data["tipo"]         = "orcamento"
            ctx.user_data["modo_revisao"] = pfm_codigo
            texto_resumo, markup = _resumo_gerar(int(doc_id))
            try:
                await query.edit_message_text(texto_resumo, reply_markup=markup, parse_mode="HTML")
            except Exception:
                await query.answer("Tela de revisão já está aberta.")

        elif acao == "pfm_ver":
            _, doc_id, pfm_codigo = partes
            await query.answer()
            html      = _gerar_html_pc(int(doc_id))
            pdf_bytes = await _html_para_pdf(html)
            await ctx.bot.send_document(
                chat_id=DONO_ID,
                document=pdf_bytes,
                filename=f"{pfm_codigo}.pdf",
                caption=f"📄 {pfm_codigo}"
            )

        elif acao == "pfm_orc":
            _, doc_id, pfm_codigo = partes
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute("SELECT caminho FROM documentos WHERE id=?", (int(doc_id),)).fetchone()
            caminho = row[0] if row else None
            if not caminho or not Path(caminho).exists():
                await query.answer("Orçamento original não encontrado.", show_alert=True)
                return
            path = Path(caminho)
            await query.answer()
            dados = path.read_bytes()
            ext = path.suffix.lower()
            if ext in (".jpg", ".jpeg", ".png", ".webp"):
                await ctx.bot.send_photo(chat_id=query.message.chat_id, photo=dados)
            else:
                await ctx.bot.send_document(chat_id=query.message.chat_id, document=dados, filename=path.name)

        elif acao in ("pfm_nfe", "pfm_comp"):
            _, doc_id_arquivo, pfm_codigo = partes
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute("SELECT caminho FROM documentos WHERE id=?", (int(doc_id_arquivo),)).fetchone()
            caminho = row[0] if row else None
            label = {"pfm_nfe": "NF-e", "pfm_comp": "comprovante"}.get(acao, acao)
            if not caminho or not Path(caminho).exists():
                await query.answer(f"Arquivo de {label} não encontrado.", show_alert=True)
                return
            path = Path(caminho)
            await query.answer()
            dados = path.read_bytes()
            ext = path.suffix.lower()
            if ext in (".jpg", ".jpeg", ".png", ".webp"):
                await ctx.bot.send_photo(chat_id=query.message.chat_id, photo=dados)
            else:
                await ctx.bot.send_document(chat_id=query.message.chat_id, document=dados, filename=path.name)

        elif acao == "pfm_entregue":
            _, doc_id, pfm_codigo = partes
            ctx.user_data["entrega_pfm_codigo"] = pfm_codigo
            ctx.user_data["entrega_doc_id_foto"] = None
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute("SELECT fornecedor FROM lancamentos WHERE pfm_codigo=?", (pfm_codigo,)).fetchone()
            forn = row[0] if row else pfm_codigo
            await query.edit_message_text(
                _texto_obs_entrega(pfm_codigo, forn),
                reply_markup=_teclado_obs_com_cancelar()
            )

        elif acao == "entrega_sel":
            _, pfm_codigo = partes
            ctx.user_data["entrega_pfm_codigo"] = pfm_codigo
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute("SELECT fornecedor FROM lancamentos WHERE pfm_codigo=?", (pfm_codigo,)).fetchone()
            forn = row[0] if row else pfm_codigo
            await query.edit_message_text(
                _texto_obs_entrega(pfm_codigo, forn),
                reply_markup=_teclado_obs_com_cancelar()
            )

        elif acao == "entrega_obs":
            _, chave = partes
            OBS_MAP = {
                "completa":  "Entrega completa",
                "parcial":   "Entrega parcial — aguardando restante",
                "avaria":    "Material com avaria",
                "diferente": "Produto diferente do pedido",
            }
            if chave == "outro":
                ctx.user_data["aguardando"] = "obs_entrega_texto"
                await query.edit_message_text("Descreva a observação:")
                return
            obs = OBS_MAP.get(chave, chave)
            pfm_codigo = ctx.user_data.pop("entrega_pfm_codigo", None)
            doc_id_foto = ctx.user_data.pop("entrega_doc_id_foto", None)
            legenda_foto = ctx.user_data.pop("entrega_legenda_foto", None)
            if not pfm_codigo:
                await query.answer("Pedido não encontrado. Tente novamente.", show_alert=True)
                return
            _salvar_entrega_db(pfm_codigo, obs)
            if doc_id_foto:
                _adicionar_foto_entrega(pfm_codigo, doc_id_foto, legenda_foto)
            texto, markup = _tela_apos_entrega(pfm_codigo)
            if texto:
                await query.edit_message_text(texto, reply_markup=markup)
            else:
                await query.edit_message_text("Entrega registrada.")

        elif acao == "entrega_cancelar":
            ctx.user_data.pop("entrega_pfm_codigo", None)
            ctx.user_data.pop("entrega_doc_id_foto", None)
            ctx.user_data.pop("entrega_legenda_foto", None)
            await query.edit_message_text(
                mostrar_boas_vindas(),
                reply_markup=teclado_boas_vindas()
            )

        elif acao == "entrega_foto_primeiro":
            ctx.user_data["aguardando"] = "foto_entrega_obs"
            await query.edit_message_text("Envie a foto ou documento da entrega:")

        elif acao == "entrega_editar":
            pfm_codigo = partes[1]
            forn, obs_ent, qtd_fotos = _buscar_estado_entrega(pfm_codigo)
            await query.edit_message_text(
                _texto_gerir_entrega(pfm_codigo, forn, obs_ent, qtd_fotos),
                reply_markup=_teclado_gerir_entrega(pfm_codigo, qtd_fotos)
            )

        elif acao == "entrega_mudar_obs":
            pfm_codigo = partes[1]
            ctx.user_data["entrega_pfm_codigo"] = pfm_codigo
            _, obs_atual, _ = _buscar_estado_entrega(pfm_codigo)
            await query.edit_message_text(
                f"#{pfm_codigo} — Mudar observação\n\nAtual: {obs_atual or '—'}",
                reply_markup=_teclado_mudar_obs(pfm_codigo)
            )

        elif acao == "entrega_editobs":
            chave = partes[1]
            pfm_codigo = ctx.user_data.get("entrega_pfm_codigo")
            if not pfm_codigo:
                await query.answer("Sessão expirada. Abra o pedido novamente.", show_alert=True)
                return
            OBS_MAP_EDIT = {
                "completa":  "Entrega completa",
                "parcial":   "Entrega parcial — aguardando restante",
                "avaria":    "Material com avaria",
                "diferente": "Produto diferente do pedido",
            }
            if chave == "outro":
                ctx.user_data["aguardando"] = "edit_obs_entrega_texto"
                await query.edit_message_text("Descreva a observação:")
                return
            obs = OBS_MAP_EDIT.get(chave, chave)
            _atualizar_obs_entrega(pfm_codigo, obs)
            forn, obs_ent, qtd_fotos = _buscar_estado_entrega(pfm_codigo)
            await query.edit_message_text(
                _texto_gerir_entrega(pfm_codigo, forn, obs_ent, qtd_fotos),
                reply_markup=_teclado_gerir_entrega(pfm_codigo, qtd_fotos)
            )

        elif acao == "entrega_trocar_foto":
            pfm_codigo = partes[1]
            ctx.user_data["aguardando"] = "foto_entrega_troca"
            ctx.user_data["entrega_pfm_codigo"] = pfm_codigo
            await query.edit_message_text("Envie a foto ou documento da entrega:")

        elif acao == "entrega_ver_fotos":
            pfm_codigo = partes[1]
            fotos = _listar_fotos_entrega(pfm_codigo)
            if not fotos:
                await query.answer("Nenhuma foto registrada.", show_alert=True)
                return
            botoes = []
            for foto_id, _doc_id_foto, legenda, caminho in fotos:
                rotulo = (legenda or "Sem legenda")[:40]
                icone = _icone_arquivo_entrega(caminho)
                botoes.append([InlineKeyboardButton(f"{icone} {rotulo}", callback_data=f"entrega_foto_ver:{foto_id}:{pfm_codigo}")])
            botoes.append([InlineKeyboardButton("← Voltar", callback_data=f"entrega_editar:{pfm_codigo}")])
            await query.edit_message_text(
                f"#{pfm_codigo} — Arquivos da entrega ({len(fotos)})",
                reply_markup=InlineKeyboardMarkup(botoes)
            )

        elif acao == "entrega_foto_ver":
            _, foto_id, pfm_codigo = partes
            with sqlite3.connect(DB_PATH) as con:
                row = con.execute(
                    "SELECT ef.legenda, d.caminho FROM entrega_fotos ef "
                    "JOIN documentos d ON d.id = ef.doc_id WHERE ef.id=?",
                    (int(foto_id),)
                ).fetchone()
            if not row or not row[1] or not Path(row[1]).exists():
                await query.answer("Foto não encontrada.", show_alert=True)
                return
            legenda, caminho = row
            await query.answer()
            path = Path(caminho)
            dados = path.read_bytes()
            ext = path.suffix.lower()
            if ext in (".jpg", ".jpeg", ".png", ".webp"):
                await ctx.bot.send_photo(chat_id=query.message.chat_id, photo=dados, caption=legenda or None)
            else:
                await ctx.bot.send_document(chat_id=query.message.chat_id, document=dados, filename=path.name, caption=legenda or None)

        elif acao == "entrega_remover_foto":
            pfm_codigo = partes[1]
            fotos = _listar_fotos_entrega(pfm_codigo)
            if not fotos:
                await query.answer("Nenhuma foto registrada.", show_alert=True)
                return
            botoes = []
            for foto_id, _doc_id_foto, legenda, _caminho in fotos:
                rotulo = (legenda or "Sem legenda")[:40]
                botoes.append([InlineKeyboardButton(f"🗑 {rotulo}", callback_data=f"entrega_foto_apagar:{foto_id}:{pfm_codigo}")])
            botoes.append([InlineKeyboardButton("← Voltar", callback_data=f"entrega_editar:{pfm_codigo}")])
            await query.edit_message_text(
                f"#{pfm_codigo} — Remover qual arquivo?",
                reply_markup=InlineKeyboardMarkup(botoes)
            )

        elif acao == "entrega_foto_apagar":
            _, foto_id, pfm_codigo = partes
            _apagar_foto_entrega(int(foto_id))
            forn, obs_ent, qtd_fotos = _buscar_estado_entrega(pfm_codigo)
            await query.edit_message_text(
                _texto_gerir_entrega(pfm_codigo, forn, obs_ent, qtd_fotos),
                reply_markup=_teclado_gerir_entrega(pfm_codigo, qtd_fotos)
            )

        elif acao == "entrega_apagar":
            pfm_codigo = partes[1]
            _apagar_entrega_db(pfm_codigo)
            texto, markup = _tela_apos_entrega(pfm_codigo)
            if texto:
                await query.edit_message_text(texto, reply_markup=markup)
            else:
                await query.edit_message_text("Entrega apagada.")

        elif acao == "entrega_voltar":
            pfm_codigo = partes[1]
            texto, markup = _tela_apos_entrega(pfm_codigo)
            if texto:
                await query.edit_message_text(texto, reply_markup=markup)
            else:
                await query.edit_message_text("Pedido não encontrado.")

        elif acao == "nfe_confirmar":
            _, doc_id_nfe, pfm_codigo = partes
            ok = vincular_nfe(pfm_codigo, int(doc_id_nfe), DB_PATH)
            if ok:
                with sqlite3.connect(DB_PATH) as con:
                    row_nfe = con.execute(
                        "SELECT caminho, dados_claude FROM documentos WHERE id=?", (int(doc_id_nfe),)
                    ).fetchone()
                if row_nfe:
                    caminho_nfe, dados_nfe = row_nfe
                    numero_nfe = _campo(dados_nfe, "Número da NF")
                    data_nfe   = _campo(dados_nfe, "Data de emissão")
                    sufixo_nfe = f"NFe {numero_nfe}" if numero_nfe != "A PREENCHER" else "NFe"
                    _arquivar_doc_financeiro(pfm_codigo, sufixo_nfe, caminho_nfe, data_nfe)
                await query.edit_message_text(
                    f"🟢 #{pfm_codigo} — NF-e vinculada. Ciclo fechado."
                )
            else:
                await query.edit_message_text(
                    f"Não foi possível vincular a NF-e ao Pedido #{pfm_codigo}.\n"
                    "O pedido pode já ter uma NF-e vinculada."
                )

        elif acao == "nfe_cancelar":
            await query.edit_message_text("NF-e não vinculada.")

        elif acao == "obra_pedidos":
            codigo  = partes[1]
            pedidos = _pedidos_obra(codigo)
            await query.edit_message_text(
                mostrar_lista_pedidos(codigo, pedidos),
                reply_markup=teclado_lista_pedidos(codigo, pedidos)
            )

        elif acao == "pedido_abrir":
            pfm_codigo = partes[1]
            pedido = buscar_pedido(pfm_codigo)
            if pedido:
                preparar_visualizacao_pedido(pedido)
                await query.edit_message_text(
                    mostrar_pedido(pedido),
                    reply_markup=teclado_pedido(pedido.doc_id, pfm_codigo, pedido.doc_id_nfe, pedido.doc_id_comprovante, pedido.qtd_fotos_entrega, pedido.obs_entrega)
                )
            else:
                await query.answer(f"Pedido {pfm_codigo} não encontrado.", show_alert=True)

        elif acao == "obra_ver":
            codigo = partes[1]
            obra = buscar_obra(codigo)
            if obra:
                await query.edit_message_text(
                    mostrar_cockpit_obra(obra),
                    reply_markup=teclado_obra(codigo, _pedidos_obra(codigo))
                )
            else:
                await query.edit_message_text(f"Obra {codigo} não encontrada.")

        elif acao == "obra_editar":
            codigo = partes[1]
            obra = buscar_obra(codigo)
            if obra:
                await query.edit_message_text(
                    f"Qual campo deseja editar em {codigo}?",
                    reply_markup=teclado_obra_campos(codigo)
                )
            else:
                await query.edit_message_text(f"Obra {codigo} não encontrada.")

        elif acao == "obra_campo":
            _, codigo, campo = partes
            ctx.user_data["aguardando"]   = f"obra_edit_{campo}"
            ctx.user_data["obra_codigo"]  = codigo
            labels = {
                "descricao":         "Descrição",
                "endereco_entrega":  "Endereço de entrega",
                "encarregado_nome":  "Nome do encarregado",
                "encarregado_fone":  "Telefone do encarregado",
                "responsavel_nome":  "Nome do responsável",
                "responsavel_fone":  "Telefone do responsável",
            }
            label = labels.get(campo, campo)
            await query.edit_message_text(f"Novo valor para {label}:")

        elif acao == "pfm_fechar":
            await query.edit_message_text("Fechado.")

        elif acao == "obra_fechar":
            await query.edit_message_text("Fechado.")

        elif acao == "obras_fechar":
            await query.edit_message_text("Fechado.")

        elif acao == "menu_inicio":
            await query.edit_message_text(
                mostrar_boas_vindas(),
                reply_markup=teclado_boas_vindas()
            )

        elif acao == "menu_obras":
            obras = _listar_obras()
            await query.edit_message_text(
                mostrar_lista_obras(obras),
                reply_markup=teclado_lista_obras(obras)
            )

        elif acao == "menu_ajuda":
            await query.edit_message_text(
                mostrar_ajuda(),
                reply_markup=teclado_ajuda(),
                parse_mode="HTML"
            )

        elif acao == "menu_nova_obra":
            await query.edit_message_text("Código da nova obra (ex: GGV04):")
            ctx.user_data["aguardando"] = "nova_obra_codigo"

    except Exception as e:
        await ctx.bot.send_message(chat_id=DONO_ID, text=f"Erro inesperado — tente novamente.\n{e}")

# ── Inicialização ──────────────────────────────────────────────────────────

async def _post_init(app):
    await app.bot.set_my_commands([
        BotCommand("help",      "Ações e consultas disponíveis"),
        BotCommand("obras",     "Lista de obras"),
        BotCommand("nova_obra", "Cadastrar obra nova"),
        BotCommand("entrega",   "Registrar entrega de pedido"),
    ])

init_db()
app = Application.builder().token(TOKEN).post_init(_post_init).build()
app.add_handler(CommandHandler("start",     start))
app.add_handler(CommandHandler("help",      ajuda))
app.add_handler(CommandHandler("obras",     obras_cmd))
app.add_handler(CommandHandler("nova_obra", nova_obra))
app.add_handler(CommandHandler("entrega",   entrega_cmd))
app.add_handler(MessageHandler(filters.PHOTO | filters.Document.ALL, receber_arquivo))
app.add_handler(CallbackQueryHandler(responder_botao))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, receber_texto))
app.add_handler(MessageHandler(filters.COMMAND, comando_desconhecido))
app.job_queue.run_repeating(_sincronizar_receita_pendentes, interval=6 * 60 * 60, first=120)
app.run_polling()
